#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Auto Document Writer — Interactive + Batch (Gemini 2.5 Flash)
- 최종 결정판 (2025-08-15)
- PDF/HWPX/DOCX 양식의 구조를 LLM으로 분석하여 문서 초안을 자동 생성
- 핵심 로직: '구조 분석 후 제목별 생성' + '사용자 지침' 반영으로 정확도 극대화
- 분석 실패 시, 사용자가 직접 제목을 입력하여 안정성 확보
- 최종 결과물은 DOCX 파일로 바로 저장
설치:
  pip install google-generativeai python-docx python-dotenv pymupdf
.env:
  GEMINI_API_KEY=YOUR_KEY
"""

import os
import re
import sys
import json
import zipfile
from pathlib import Path
from typing import Dict, List, Tuple, Optional

from dotenv import load_dotenv
import google.generativeai as genai

# ---- Optional imports ----
try:
    import fitz  # PyMuPDF for PDF
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument
    from docx.shared import RGBColor
except Exception:
    DocxDocument = None

# --------------------------
# Gemini 초기화
# --------------------------
def init_gemini(model_name: str):
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY 가 .env 에 없습니다.")
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(model_name)

# --------------------------
# 유틸: 파일/문서 도우미
# --------------------------
PLACEHOLDER_PATTERN = re.compile(r"\{\{([^{}]+)\}\}")

def read_json_or_text(p: Path) -> str:
    raw = p.read_text(encoding="utf-8", errors="ignore")
    try:
        obj = json.loads(raw)
        return json.dumps(obj, ensure_ascii=False, indent=2)
    except Exception:
        return raw

def detect_placeholders_in_text(text: str) -> List[str]:
    return sorted(set(m.group(1).strip() for m in PLACEHOLDER_PATTERN.finditer(text)))

def extract_text_from_pdf(p: Path, max_bytes: int = 50000) -> str:
    if fitz is None: return "(PyMuPDF 미설치) PDF 텍스트 추출 불가"
    doc = fitz.open(p.as_posix())
    pages = []
    total = 0
    for i, page in enumerate(doc, start=1):
        if total >= max_bytes: break
        try:
            t = page.get_text("text")
        except Exception:
            t = "(텍스트 추출 실패)"
        pages.append(t)
        total += len(t.encode("utf-8", errors="ignore"))
    doc.close()
    return "\n\n".join(pages)

def load_docx_and_plaintext(p: Path) -> Tuple[Optional[object], str]:
    if DocxDocument is None: return None, ""
    doc = DocxDocument(p.as_posix())
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            texts.append("\t".join(cell.text for cell in row.cells))
    return doc, "\n".join(texts)

def hwpx_read_xml(p: Path) -> Tuple[bytes, List[str]]:
    with zipfile.ZipFile(p, 'r') as z:
        texts = []
        for name in z.namelist():
            if name.lower().endswith(".xml"):
                with z.open(name) as f:
                    try:
                        texts.append(f.read().decode("utf-8", errors="ignore"))
                    except Exception: pass
        full = "\n".join(texts)
        return full.encode("utf-8"), detect_placeholders_in_text(full)

# --------------------------
# 프롬프트 빌더
# --------------------------
def build_prompt(plan_text: str, feature_json_text: Optional[str], task_desc: str, tone: str, style: str, leave_blanks: bool, writing_instructions: Optional[str] = None, must_include_content: Optional[str] = None) -> str:
    rules = f"""
당신은 정부/공모/연구개발 양식에 맞춰 문서를 작성하는 전문 보조자입니다.
[문서 규칙]
- 어투/톤: {tone}
- 문장 스타일: {style}
- 과장 금지, 사실 기반 표현.
- 불확실/정보부족: {"빈칸으로 남겨두세요" if leave_blanks else "<<UNSURE: 간단 사유/질문 >>"}
- 개인정보/허위 기재 금지.
"""
    if writing_instructions:
        rules += f"\n[사용자 제공 작성 요령]\n{writing_instructions}\n"
    
    if must_include_content:
        rules += f"\n[필수 포함 내용/키워드]\n{must_include_content}\n"

    rules += f"""
[참조 기획서/기능명세]
{plan_text}
"""
    if feature_json_text: rules += f"\n[추가 기능 명세(JSON)]\n{feature_json_text}\n"
    if task_desc: rules += f"\n[작업 설명]\n{task_desc}\n"
    return rules.strip()

# --------------------------
# LLM 호출 및 처리 로직
# --------------------------
def llm_generate(model, prompt: str, instruction: str) -> str:
    full = prompt + "\n\n[작성 요청]\n" + instruction
    gen_cfg = genai.types.GenerationConfig(temperature=0.3, top_p=0.9, max_output_tokens=2048)
    safety_settings = [{"category": c, "threshold": "BLOCK_NONE"} for c in ["HARM_CATEGORY_HARASSMENT", "HARM_CATEGORY_HATE_SPEECH", "HARM_CATEGORY_SEXUAL", "HARM_CATEGORY_DANGEROUS_CONTENT"]]
    try:
        resp = model.generate_content(full, generation_config=gen_cfg, safety_settings=safety_settings)
        return resp.text.strip()
    except Exception as e:
        print(f"[경고] LLM 호출 예외: {e}", file=sys.stderr)
        return ""

def get_headings_from_structure(model, raw_text: str) -> List[str]:
    """[개선판] 문서 텍스트(표 포함)에서 LLM을 이용해 주요 제목 목록을 추출합니다."""
    print("  - LLM으로 문서 구조를 분석하여 제목 목록을 추출합니다...")
    prompt = f"""
아래 텍스트는 문서 양식(.docx, .pdf 등)에서 추출한 내용입니다. 표(Table) 형식의 양식이 포함될 수 있습니다.
당신의 임무는 이 문서에서 내용을 작성해야 할 **핵심 항목 또는 제목 목록을 순서대로 추출**하는 것입니다.

[추출 규칙]
- 표 형식의 경우, 보통 첫 번째 열(column)에 '프로젝트명', '개발배경 및 목적'과 같은 항목 이름이 있고 두 번째 열은 비어있거나 설명이 있습니다. **첫 번째 열의 내용을 제목으로 정확하게 추출하세요.**
- 일반 텍스트 형식의 경우, 제목처럼 보이는 짧은 구절을 추출하세요.
- '#', '※', '□' 와 같은 기호나 '예시)', '...하는 방법' 같은 설명 문구는 제외하고 항목 이름 자체만 깔끔하게 추출해주세요.
- 각 제목은 한 줄에 하나씩만 출력하세요.

[분석할 문서 텍스트]
{raw_text[:20000]}

[추출된 제목 목록]
"""
    try:
        gen_cfg = genai.types.GenerationConfig(temperature=0.0, max_output_tokens=1024)
        resp = model.generate_content(prompt, generation_config=gen_cfg)
        headings = [line.strip() for line in resp.text.strip().splitlines() if line.strip()]
        cleaned_headings = []
        for h in headings:
            if h in ["구 분", "세 부 내 용", "결과물"]: continue
            cleaned_h = re.sub(r"\s*,\s*", " ", h).strip()
            cleaned_h = re.sub(r"^\d+\.\s*|[-\*]\s*|[□]\s*", "", cleaned_h).strip()
            if cleaned_h: cleaned_headings.append(cleaned_h)
        return cleaned_headings
    except Exception as e:
        print(f"[경고] 제목 목록 추출 중 예외 발생: {e}", file=sys.stderr)
        return []

def generate_content_for_headings(model, headings: List[str], plan_text: str, feature_json_text: Optional[str], tone: str, style: str, leave_blanks: bool, writing_instructions: str, must_include_content: str) -> str:
    """제목 목록을 받아 각 제목에 해당하는 내용을 생성하고 합칩니다."""
    base_prompt = build_prompt(plan_text, feature_json_text, "문서의 각 섹션을 해당 제목 아래에 작성하세요.", tone, style, leave_blanks, writing_instructions, must_include_content)
    full_content = []
    for i, heading in enumerate(headings, 1):
        print(f"  - 섹션 생성 중 ({i}/{len(headings)}): {heading}")
        instruction = f"아래 제목에 해당하는 섹션의 본문을 작성하세요. 제목은 결과에 포함하지 마세요.\n\n[제목]\n{heading}"
        body = llm_generate(model, base_prompt, instruction)
        section = f"## {heading}\n\n{body}\n\n"
        full_content.append(section)
    return "".join(full_content)

# --------------------------
# 핵심 로직: 구조-모방 모드
# --------------------------
def process_structure_mimic_mode(model, raw_text: str, tpl_path: Path, plan_text: str, feature_json_text: Optional[str], tone: str, style: str, leave_blanks: bool, out_override: str, writing_instructions: str, must_include_content: str):
    """[최종 결정판] 문서 구조에서 제목 목록을 추출하고, 실패 시 사용자 입력을 받아 안정성을 극대화합니다."""
    headings = get_headings_from_structure(model, raw_text)
    
    if not headings:
        print("[오류] 문서에서 제목 구조를 추출하는데 실패했습니다.", file=sys.stderr)
        print("\n수동으로 제목을 입력해주세요. 한 줄에 하나씩 입력하고, 완료되면 빈 줄에서 엔터를 누르세요.")
        user_headings = []
        while True:
            try:
                line = input("> ")
                if not line.strip(): break
                user_headings.append(line.strip())
            except EOFError: break
        
        headings = user_headings
        if not headings:
            print("[오류] 수동으로 입력된 제목이 없습니다. 작업을 중단합니다.", file=sys.stderr)
            return
        print(f"\n[알림] 수동으로 입력된 {len(headings)}개의 제목으로 문서 생성을 시작합니다.")

    print(f"  - 총 {len(headings)}개의 제목을 기반으로 문서 생성을 시작합니다.")
    md_all = generate_content_for_headings(model, headings, plan_text, feature_json_text, tone, style, leave_blanks, writing_instructions, must_include_content)

    if DocxDocument is None:
        print("[경고] python-docx가 설치되지 않아 DOCX 저장이 불가합니다. Markdown으로 대신 저장합니다.", file=sys.stderr)
        out_path = resolve_output_path(out_override, tpl_path.with_name(tpl_path.stem + "_completed.md"), ".md")
        out_path.write_text(md_all, encoding="utf-8")
        print(f"✅ 완성된 Markdown 문서 저장: {out_path}")
        return

    out_path_docx = resolve_output_path(out_override, tpl_path.with_name(tpl_path.stem + "_completed.docx"), ".docx")
    doc = DocxDocument()
    for block in md_all.strip().split('\n\n'):
        if not block.strip(): continue
        lines = block.split('\n')
        first_line = lines[0].strip()
        if first_line.startswith('## '):
            doc.add_heading(first_line[3:], level=2)
            if len(lines) > 1: doc.add_paragraph('\n'.join(lines[1:]))
        elif first_line.startswith('# '):
            doc.add_heading(first_line[2:], level=1)
            if len(lines) > 1: doc.add_paragraph('\n'.join(lines[1:]))
        else:
            doc.add_paragraph(block)
    
    doc.save(out_path_docx.as_posix())
    print(f"✅ 완성된 DOCX 문서 저장: {out_path_docx}")

    out_path_md = out_path_docx.with_suffix(".md")
    out_path_md.write_text(md_all, encoding="utf-8")
    print(f"  - 원본 텍스트(Markdown) 저장: {out_path_md}")

# --------------------------
# 보조 함수들
# --------------------------
def resolve_output_path(user_in: str, default_path: Path, desired_suffix: str) -> Path:
    if not user_in: return default_path
    p = Path(user_in)
    if p.is_dir(): return p / default_path.name
    if p.suffix == "": return p.with_suffix(desired_suffix)
    if not p.parent.exists():
        try: p.parent.mkdir(parents=True, exist_ok=True)
        except: pass
    return p

# --------------------------
# 인터랙티브 메인
# --------------------------
def main():
    print(" Auto Document Writer (Gemini 2.5 Flash) — 최종 결정판\n")

    plan_path_str = input(" 기획서/기능명세 파일 경로 (.json, .txt): ").strip()
    template_path_str = input(" 작성할 양식 파일 경로 (.pdf, .docx, .hwpx): ").strip()
    tone = input(" 어투/톤 (기본: 격식체): ").strip() or "격식체"
    style = input(" 문장 스타일 (기본: 개조식): ").strip() or "개조식"
    leave_blanks = input(" 불확실시 빈칸 처리? (y/N): ").strip().lower() == "y"
    
    print("\n 양식에 대한 '작성 요령'이나 특별한 지침이 있다면 아래에 붙여넣어 주세요.")
    writing_instructions = input(" 추가 작성 지침 (선택, 없으면 엔터): ").strip()

    print("\n 문서에 반드시 포함되어야 할 내용이나 키워드가 있다면 아래에 입력해주세요.")
    must_include_content = input(" 필수 포함 내용 (선택, 없으면 엔터): ").strip()

    # --- 모델 이름 'gemini-2.5-flash'로 고정 ---
    model_name = "gemini-2.5-flash"
    print(f"\n Gemini 모델: {model_name} (고정)")
    
    out_override = input(" 출력 파일/폴더 경로 (미입력시 자동 결정): ").strip()

    plan_path = Path(plan_path_str)
    tpl_path = Path(template_path_str)

    if not plan_path.exists(): print("기획서 파일을 찾을 수 없습니다.", file=sys.stderr); sys.exit(1)
    if not tpl_path.exists(): print("양식 파일을 찾을 수 없습니다.", file=sys.stderr); sys.exit(1)

    try:
        model = init_gemini(model_name)
    except Exception as e:
        print(f"Gemini 초기화 실패: {e}", file=sys.stderr); sys.exit(1)

    plan_text = read_json_or_text(plan_path)
    feature_json_text = json.dumps(json.loads(plan_text), ensure_ascii=False, indent=2) if plan_text.strip().startswith(('{', '[')) else None
    
    ext = tpl_path.suffix.lower()
    
    print(f"\n[구조-모방 모드] {ext} 파일의 구조를 분석하여 새 문서를 생성합니다.")
    raw_text = ""
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(tpl_path)
    elif ext == ".docx":
        _, raw_text = load_docx_and_plaintext(tpl_path)
    elif ext == ".hwpx":
        xml_blob, _ = hwpx_read_xml(tpl_path)
        raw_text = xml_blob.decode("utf-8", errors="ignore")
    elif ext == ".hwp":
        print("[알림] .hwp 파일의 직접적인 텍스트 추출은 지원되지 않습니다. HWPX로 변환하여 사용해주세요.", file=sys.stderr)
        raw_text = "과제 개요\n주요 내용\n기대 효과"
    else:
        print(f"지원하지 않는 확장자: {ext}", file=sys.stderr); sys.exit(2)

    process_structure_mimic_mode(model, raw_text, tpl_path, plan_text, feature_json_text, tone, style, leave_blanks, out_override, writing_instructions, must_include_content)

if __name__ == "__main__":
    main()