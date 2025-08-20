from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from django.contrib.auth import authenticate, get_user_model
from rest_framework_simplejwt.tokens import RefreshToken

from .serializers import UserSerializer

User = get_user_model()  # People 모델 불러오기

from .serializers import UserSerializer, SignupSerializer
#회원 가입
class SignupView(APIView):
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        serializer = SignupSerializer(data=request.data)
        if serializer.is_valid():
            user = serializer.save()
            return Response({"message": "회원가입 성공"}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
# 이메일 중복 확인
class EmailCheckView(APIView):
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        email = request.data.get("email")
        if not email:
            return Response({"error": "이메일을 입력해주세요."}, status=status.HTTP_400_BAD_REQUEST)

        if User.objects.filter(email=email).exists():
            return Response({"is_duplicate": True, "message": "이미 사용 중인 이메일입니다."}, status=status.HTTP_200_OK)
        else:
            return Response({"is_duplicate": False, "message": "사용 가능한 이메일입니다."}, status=status.HTTP_200_OK)



class LoginView(APIView):
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        email = request.data.get('email')
        password = request.data.get('password')

        # ✅ email 기반 로그인 대응
        user = User.objects.filter(email=email).first()
        if user and user.check_password(password):
            refresh = RefreshToken.for_user(user)
            return Response({
                'access': str(refresh.access_token),
                'refresh': str(refresh),
                'user_id': user.user_id,
                'username': user.username
            })

        return Response({"error": "이메일 또는 비밀번호가 틀렸습니다."},
                        status=status.HTTP_401_UNAUTHORIZED)


from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from .serializers import UserSerializer

class UserInfoView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = User.objects.get(user_id=request.user.user_id)  # 🔥 핵심 수정
        serializer = UserSerializer(user)
        return Response(serializer.data)

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from .models import Project, RequirementDraft, Requirement, SimilarProject, TeamMember, ProjectTimeline, OutputDocument
from .serializers import ProjectSerializer
from .gemini_parserv2 import generate_feature_list

class ProjectCreateView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        title = request.data.get("title")
        description = request.data.get("description")

        if not title or not description:
            return Response({"error": "제목과 설명은 필수입니다."}, status=status.HTTP_400_BAD_REQUEST)

        project = Project.objects.create(
            title=title,
            description=description,
            user=request.user
        )

        serializer = ProjectSerializer(project)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
# --- ADD: register-from-file (프로젝트 생성만; 초안 생성은 별도 Gemini1GenerateView에서 실행) ---
# --- register-from-file (프로젝트 생성만; 초안 생성은 별도) ---
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
import json, os, io, csv

from .models import Project

def _read_any_file(upload):
    name = getattr(upload, "name", "upload")
    ext = os.path.splitext(name)[1].lower()
    buf = upload.read()
    source_label = f"file:{name}"
    try:
        if ext in [".txt", ".md", ".markdown", ".log", ".ini", ".conf"]:
            text = buf.decode("utf-8", errors="ignore")
            return {"plan_text": text, "source_label": source_label}
        if ext == ".csv":
            try:
                data = buf.decode("utf-8", errors="ignore").splitlines()
                rdr = csv.reader(data)
                lines = []
                for row in rdr:
                    lines.append(", ".join(str(c) for c in row))
                return {"plan_text": "\n".join(lines), "source_label": source_label}
            except Exception as e:
                raise ValueError(f"CSV 파싱 실패: {e}")
        if ext == ".json":
            try:
                obj = json.loads(buf.decode("utf-8", errors="ignore"))
            except Exception as e:
                raise ValueError(f"JSON 파싱 실패: {e}")
            plan_text = ""
            plan_lines = obj.get("기획서원문") or obj.get("plan_lines") or obj.get("plan") or obj.get("description")
            if isinstance(plan_lines, list):
                plan_text = "\n".join(str(x) for x in plan_lines)
            elif isinstance(plan_lines, str):
                plan_text = plan_lines
            if not plan_text:
                plan_text = json.dumps(obj, ensure_ascii=False, indent=2)
            return {"plan_text": plan_text or "", "source_label": source_label}
        if ext in [".yml", ".yaml"]:
            try:
                import yaml
            except Exception:
                raise ValueError("YAML을 읽으려면 pyyaml가 필요합니다. (pip install pyyaml)")
            try:
                obj = yaml.safe_load(buf.decode("utf-8", errors="ignore"))
            except Exception as e:
                raise ValueError(f"YAML 파싱 실패: {e}")
            text = "" if obj is None else yaml.safe_dump(obj, allow_unicode=True, sort_keys=False)
            return {"plan_text": text, "source_label": source_label}
        if ext == ".doc":
            try:
                import textract
                text = textract.process(io.BytesIO(buf)).decode("utf-8", errors="ignore")
            except Exception:
                raise ValueError("DOC 읽기엔 textract 및 시스템 도구가 필요합니다. 가능하면 DOCX로 변환 권장.")
            return {"plan_text": text, "source_label": source_label}
        if ext == ".docx":
            try:
                from docx import Document
            except Exception:
                raise ValueError("DOCX를 읽으려면 python-docx가 필요합니다. (pip install python-docx)")
            doc = Document(io.BytesIO(buf))
            text = "\n".join(p.text for p in doc.paragraphs)
            return {"plan_text": text, "source_label": source_label}
        if ext == ".pdf":
            try:
                import pdfplumber
            except Exception:
                raise ValueError("PDF를 읽으려면 pdfplumber가 필요합니다. (pip install pdfplumber)")
            texts = []
            with pdfplumber.open(io.BytesIO(buf)) as pdf:
                for pg in pdf.pages:
                    texts.append(pg.extract_text() or "")
            return {"plan_text": "\n".join(texts), "source_label": source_label}
        raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")
    finally:
        try:
            upload.close()
        except Exception:
            pass


class ProjectRegisterFromFileView(APIView):
    """
    POST /api/project/register-from-file/
    multipart/form-data:
      - title: str (필수)
      - file : 업로드 파일 (필수)
    동작:
      1) 파일에서 원문(plan_text)만 추출
      2) Project(title, description=plan_text) 생성
      3) (초안 생성은 하지 않음) → 이후 /api/project/<project_id>/generate-gemini1/ 호출
    """
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser, JSONParser]

    def post(self, request):
        title = (request.data.get("title") or "").strip()
        upload = request.FILES.get("file") or request.FILES.get("plan_file")

        if not title:
            return Response({"ok": False, "error": "title은 필수입니다."}, status=status.HTTP_400_BAD_REQUEST)
        if not upload:
            return Response({"ok": False, "error": "파일이 없습니다. 'file' 필드로 업로드하세요."}, status=status.HTTP_400_BAD_REQUEST)

        # 1) 파일에서 원문 추출
        try:
            parsed = _read_any_file(upload)
            plan_text = (parsed.get("plan_text") or "").strip()
            used_source = parsed.get("source_label") or "file:upload"
            if not plan_text:
                return Response({"ok": False, "error": "파일에서 원문을 추출하지 못했습니다."}, status=status.HTTP_400_BAD_REQUEST)
            MAX_CHARS = 20000
            if len(plan_text) > MAX_CHARS:
                plan_text = plan_text[:MAX_CHARS]
        except ValueError as ve:
            return Response({"ok": False, "error": str(ve)}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"ok": False, "error": f"업로드 처리 중 오류: {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # 2) Project 생성 (description 동기화)
        try:
            project = Project.objects.create(
                user=request.user,
                title=title[:200],
                description=plan_text,
            )
        except Exception as e:
            return Response({"ok": False, "error": f"Project 생성 실패: {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # 3) 응답 (호환 모드: 평탄화 + 중첩 + ok 플래그, 200으로 반환)
        payload = {
            "ok": True,
            "message": "프로젝트가 생성되었습니다. (초안 생성은 별도 호출)",
            # 평탄화(프론트가 최상위에서 읽는 경우 대비)
            "project_id": project.project_id,
            "projectId": project.project_id,  # camelCase도 함께 제공 (프론트 호환)
            "title": project.title,
            "description": project.description,
            "used_source": used_source,
            "next": f"/api/project/{project.project_id}/generate-gemini1/",
            # 중첩 블록(기존 새 응답 형식 유지)
            "project": {
                "project_id": project.project_id,
                "title": project.title,
                "description": project.description,
                "created_at": project.created_at,
            },
        }
        return Response(payload, status=status.HTTP_200_OK)




from django.conf import settings
from django.utils import timezone
from django.shortcuts import get_object_or_404
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
import json, os, io
import time # ✅ time 모듈 import

from .models import Project, RequirementDraft
from .gemini_parserv2 import generate_feature_list, export_tabular_files

class Gemini1GenerateView(APIView):
    """
    POST /api/project/<project_id>/generate-gemini1/
    (설명은 기존과 동일)
    """
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser, JSONParser]

    def _read_uploaded_file(self, upload):
        # (이 함수는 변경 없음)
        name = getattr(upload, "name", "upload")
        ext = os.path.splitext(name)[1].lower()
        buf = upload.read()

        try:
            if ext in [".txt", ".md"]:
                return buf.decode("utf-8", errors="ignore")
            if ext == ".docx":
                try:
                    from docx import Document
                except Exception:
                    raise ValueError("DOCX를 읽으려면 python-docx가 필요합니다. (pip install python-docx)")
                doc = Document(io.BytesIO(buf))
                return "\n".join(p.text for p in doc.paragraphs)
            if ext == ".pdf":
                try:
                    import pdfplumber
                except Exception:
                    raise ValueError("PDF를 읽으려면 pdfplumber가 필요합니다. (pip install pdfplumber)")
                texts = []
                with pdfplumber.open(io.BytesIO(buf)) as pdf:
                    for pg in pdf.pages:
                        texts.append(pg.extract_text() or "")
                return "\n".join(texts)
            raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")
        finally:
            try:
                upload.close()
            except Exception:
                pass

    def post(self, request, project_id):
        # 0) 프로젝트 검증
        project = get_object_or_404(Project, project_id=project_id, user=request.user)

        # 1) 입력 원문 확보
        # (이 부분은 변경 없음)
        used_source = "project.description"
        try:
            plan_text = ""
            upload = request.FILES.get("file") or request.FILES.get("plan_file")
            if upload:
                plan_text = self._read_uploaded_file(upload)
                used_source = f"file:{getattr(upload, 'name', 'upload')}"
            else:
                plan_text = (request.data.get("plan_text", "") or "").strip()
                if plan_text:
                    used_source = "plan_text"
                else:
                    plan_text = (project.description or "").strip()
            if not plan_text:
                return Response({"error": "입력 텍스트가 비어 있습니다."}, status=status.HTTP_400_BAD_REQUEST)
            MAX_CHARS = 20000
            if len(plan_text) > MAX_CHARS:
                plan_text = plan_text[:MAX_CHARS]
        except ValueError as ve:
            return Response({"error": str(ve)}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": f"업로드 처리 중 오류: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # 1-1) 원문 동기화
        # (이 부분은 변경 없음)
        try:
            if used_source != "project.description" and plan_text and plan_text != (project.description or ""):
                project.description = plan_text
                project.save(update_fields=["description"])
        except Exception:
            pass

        # ======================================================================
        # ✅ 2) 기능 리스트 생성 (안정성 강화)
        # ======================================================================
        final_features = []
        # ✅ 기본 반복 횟수를 3으로 조정 (환경 변수로 덮어쓰기 가능)
        MAX_PASSES = int(os.getenv("AUTO_PASSES", "3"))

        for pass_count in range(1, MAX_PASSES + 1):
            # ✅ 두 번째 호출부터는 1초 지연 시간을 두어 API 과부하 방지
            if pass_count > 1:
                time.sleep(1)

            print(f"-> Django View: 기능 추출 패스 #{pass_count} 진행")
            try:
                new_features = generate_feature_list(plan_text, existing_features=final_features)
                if new_features:
                    final_features.extend(new_features)
                    print(f"   (패스 #{pass_count}: {len(new_features)}개 기능 추가)")
                else:
                    print(f"   (패스 #{pass_count}: 새로운 기능 없음, 반복 종료)")
                    break
            except Exception as e:
                return Response({
                    "error": f"Gemini1 처리 중 패스 #{pass_count}에서 오류가 발생했습니다.",
                    "detail": str(e)
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        # ======================================================================

        # 3) 드래프트 저장
        # (이 부분은 변경 없음, final_features 사용)
        try:
            first = final_features[0] if final_features else {}
            feature_name = ((first.get("기능명") or first.get("feature_name") or "").strip() or "기능명 없음")
            summary = ""
            desc = first.get("기능설명") or {}
            if isinstance(desc, dict):
                summary = (desc.get("목적") or desc.get("핵심역할") or "").strip()
            if not summary:
                outputs = first.get("출력값") or {}
                if isinstance(outputs, dict):
                    summary = (outputs.get("요약정보") or "").strip()
            draft = RequirementDraft.objects.create(
                project=project,
                source="gemini_1",
                content=json.dumps(final_features, ensure_ascii=False),
                generated_by="gemini_1",
                feature_name=feature_name,
                summary=summary or "설명 없음",
                score_by_model=0.0,
            )
        except Exception as e:
            return Response({"error": f"RequirementDraft 저장 실패: {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # 3-1) 사용자용 파일 생성
        # (이 부분은 변경 없음, final_features 사용)
        warnings = []
        try:
            media_root = getattr(settings, "MEDIA_ROOT", os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media"))
            media_url  = getattr(settings, "MEDIA_URL", "/media/")
            out_dir = os.path.join(media_root, "drafts")
            os.makedirs(out_dir, exist_ok=True)
            ts = timezone.now().strftime("%Y%m%d_%H%M%S")
            base_name = f"project{project.project_id}_{ts}_g1"
            base_path = os.path.join(out_dir, base_name)
            plan_lines = (project.description or "").splitlines()
            json_path = f"{base_path}.json"
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump({"기획서원문": plan_lines, "기능목록": final_features}, f, ensure_ascii=False, indent=2)
            json_url = f"{media_url}drafts/{os.path.basename(json_path)}"
            xlsx_url = None
            try:
                export_tabular_files(plan_lines, final_features, base_path)
                xlsx_url = f"{media_url}drafts/{os.path.basename(base_path)}.xlsx"
            except (ModuleNotFoundError, ImportError):
                warnings.append("pandas/openpyxl이 없어 G1 XLSX 생성을 건너뜁니다.")
            except Exception as e:
                warnings.append(f"G1 엑셀 생성 실패: {e}")
        except Exception as e:
            return Response({"error": f"G1 파일 저장 실패: {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # 4) 응답
        # (이 부분은 변경 없음, final_features 사용)
        return Response(
            {
                "message": "Gemini 1 기능 초안 생성 완료",
                "draft_id": draft.RequirementDraft_id,
                "features": final_features,
                "score_by_model": 0.0,
                "used_source": used_source,
                "files": {"json": json_url, "xlsx": xlsx_url},
                "warnings": warnings or None,
            },
            status=status.HTTP_201_CREATED,
        )


    

# 제미나이 2 검증용
# --- Gemini2: Refine & Excel 동기화 (교체본) ---
from django.conf import settings
from django.shortcuts import get_object_or_404
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
import os, json, re
from django.utils import timezone

import google.generativeai as genai
from google.generativeai.types import GenerationConfig

from .models import Project, RequirementDraft
from .gemini_refiner import make_refine_prompt, export_excel_from_features

class Gemini2RefineView(APIView):
    """
    POST /api/project/<project_id>/refine-gemini2/
    body:
      { "draft_id": 123 }   # (옵션) 없으면 최신 gemini_1 초안 사용
    동작:
      1) gemini_1 초안(JSON 배열) 로드 (파일 입력은 받지 않음)
      2) make_refine_prompt(plan_text, features)로 프롬프트 생성
      3) gemini-2.5-flash로 정제 JSON 생성
      4) JSON 파일 + (가능하면) 엑셀 동기화 파일을 MEDIA_ROOT에 저장
      5) gemini_2 초안으로 RequirementDraft 저장
      6) refined_content + 파일 URL 반환 (xlsx는 미설치 시 None)
    """
    permission_classes = [permissions.IsAuthenticated]

    def _strip_code_fence(self, s: str) -> str:
        s = (s or "").strip()
        # ```json ... ``` 제거
        if s.startswith("```"):
            s = re.sub(r"^```(?:json)?", "", s, flags=re.I).strip()
            s = re.sub(r"```$", "", s).strip()
        return s

    def post(self, request, project_id):
        # ❌ 파일 입력 금지: G2는 검증/정제 전용
        if request.FILES:
            return Response(
                {"error": "파일 업로드는 Gemini 1에서만 지원됩니다."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # 0) 프로젝트 확인
        project = get_object_or_404(Project, project_id=project_id, user=request.user)

        # 1) 입력 초안 선택 (지정 없으면 최신 gemini_1)
        draft_id = request.data.get("draft_id")
        if draft_id:
            src_draft = get_object_or_404(
                RequirementDraft, pk=draft_id, project=project, source="gemini_1"
            )
        else:
            src_draft = (
                RequirementDraft.objects
                .filter(project=project, source="gemini_1")
                .order_by("-created_at", "-RequirementDraft_id")
                .first()
            )
            if not src_draft:
                return Response(
                    {"error": "Gemini 1 기능 초안이 없습니다. 먼저 G1을 실행하세요."},
                    status=status.HTTP_404_NOT_FOUND
                )

        # 2) 초안 content(JSON 문자열) → list
        try:
            features = json.loads(src_draft.content or "[]")
            if not isinstance(features, list) or not features:
                return Response(
                    {"error": "초안 content는 비어있지 않은 JSON 배열이어야 합니다."},
                    status=status.HTTP_400_BAD_REQUEST
                )
        except Exception as e:
            return Response(
                {"error": f"초안 JSON 파싱 실패: {e}"},
                status=status.HTTP_400_BAD_REQUEST
            )

        # 원문 텍스트 (너무 길면 잘라서 프롬프트 보호)
        plan_text = (project.description or "").strip()
        if len(plan_text) > 20000:
            plan_text = plan_text[:20000]

        # 3) Refiner 프롬프트 생성 및 LLM 호출
        try:
            # 안전한 구성: settings 또는 환경변수에서 키가 있으면 설정
            api_key = getattr(settings, "GEMINI_API_KEY_2", None) or os.getenv("GEMINI_API_KEY_2")
            if api_key:
                genai.configure(api_key=api_key)

            prompt = make_refine_prompt(plan_text, features)
            model = genai.GenerativeModel("gemini-2.5-flash")
            resp = model.generate_content(
                prompt,
                generation_config=GenerationConfig(
                    temperature=0.2,
                    response_mime_type="application/json"  # ✅ JSON 강제
                )
            )
            text = self._strip_code_fence(getattr(resp, "text", "") or "")

            try:
                refined = json.loads(text)
            except json.JSONDecodeError:
                # 파싱 실패 시 원문 텍스트 그대로 보존
                refined = text

        except Exception as e:
            msg = str(e)
            low = msg.lower()
            if "429" in msg or "rate" in low or "quota" in low or "exceeded" in low:
                return Response({
                    "error": "요청이 많아 일시적으로 제한되었습니다(HTTP 429).",
                    "hint": "토큰 만료/쿼터 초과 가능. .env의 GEMINI_API_KEY_2 확인/갱신 후 서버 재시작.",
                    "detail": msg
                }, status=status.HTTP_429_TOO_MANY_REQUESTS)
            if "401" in msg or "unauthorized" in low or "invalid" in low or "permission" in low:
                return Response({
                    "error": "인증 실패(HTTP 401).",
                    "hint": "잘못된 키일 수 있습니다. .env의 GEMINI_API_KEY_2 확인/교체 후 재시작.",
                    "detail": msg
                }, status=status.HTTP_401_UNAUTHORIZED)
            return Response(
                {"error": "Gemini 2 처리 중 오류가 발생했습니다.", "detail": msg},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

        # 4) 파일 저장(MEDIA_ROOT/refine/) - XLSX는 의존성 없으면 건너뜀
        warnings = []
        media_root = getattr(settings, "MEDIA_ROOT", os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media"))
        media_url  = getattr(settings, "MEDIA_URL", "/media/")
        out_dir = os.path.join(media_root, "refine")
        os.makedirs(out_dir, exist_ok=True)

        ts = timezone.now().strftime("%Y%m%d_%H%M%S")
        base = f"project{project.project_id}_{ts}_g2"  # ✅ g2 접미사
        json_path = os.path.join(out_dir, f"{base}_fix.json")
        xlsx_path = os.path.join(out_dir, f"{base}_fix.xlsx")

        # JSON 저장 ({"정제기획서": refined} 래핑) — 항상 생성
        try:
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump({"정제기획서": refined}, f, ensure_ascii=False, indent=2)
            json_url = f"{media_url}refine/{os.path.basename(json_path)}"
        except Exception as e:
            return Response({"error": f"JSON 파일 저장 실패: {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # 엑셀 동기화: ⚠️ G2는 가능한 한 '정제 결과' 기준으로 만듦
        # refined가 list면 그걸 사용, 아니면 G1 features로 폴백
        xlsx_url = None
        try:
            plan_lines = (project.description or "").splitlines()
            refined_features_for_file = refined if isinstance(refined, list) else features
            export_excel_from_features(plan_lines, refined_features_for_file, xlsx_path)
            xlsx_url = f"{media_url}refine/{os.path.basename(xlsx_path)}"
        except (ModuleNotFoundError, ImportError):
            warnings.append("pandas/openpyxl이 없어 XLSX 생성을 건너뜁니다. 'pip install pandas openpyxl' 후 재시도하세요.")
        except Exception as e:
            warnings.append(f"엑셀 생성 실패: {e}")

        # 5) 정제 결과를 gemini_2 초안으로 저장
        try:
            refined_text_for_db = (
                json.dumps(refined, ensure_ascii=False, indent=2)
                if not isinstance(refined, str) else refined
            )
            new_draft = RequirementDraft.objects.create(
                project=project,
                source="gemini_2",
                content=refined_text_for_db,
                generated_by="gemini_2",
                feature_name="Refined Draft by Gemini 2",
                summary="Gemini 2가 정제한 결과(JSON/Excel 동기화).",
                score_by_model=0.0,
            )
        except Exception as e:
            return Response(
                {"error": f"RequirementDraft 저장 실패: {e}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

        # 6) 응답
        return Response({
            "message": "Gemini 2 정제 완료 (JSON/Excel 동기화)",
            "draft_id": new_draft.RequirementDraft_id,
            "refined_content": refined,      # JSON or str
            "files": {
                "json": json_url,
                "xlsx": xlsx_url  # 의존성 없으면 None
            },
            "warnings": warnings or None
        }, status=status.HTTP_201_CREATED)
    
# views.py
class ProjectG2FilesView(APIView):
    """
    GET /api/project/<project_id>/files/g2/
    refine/ 밑에 저장된 G2(JSON/XLSX) 파일 목록을 최신순으로 반환
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, project_id):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)

        media_root = getattr(settings, "MEDIA_ROOT",
                             os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media"))
        media_url  = getattr(settings, "MEDIA_URL", "/media/")
        pid = project.project_id

        refine_dir = os.path.join(media_root, "refine")
        json_paths = sorted(glob(os.path.join(refine_dir, f"project{pid}_*_g2_fix.json")), reverse=True)
        xlsx_set   = set(glob(os.path.join(refine_dir, f"project{pid}_*_g2_fix.xlsx")))

        def _ts_to_iso(fname_base: str):
            # project{pid}_YYYYMMDD_HHMMSS_g2_fix
            name = os.path.basename(fname_base)
            ts = name.split(f"project{pid}_")[1].split("_g2", 1)[0]
            try:
                return datetime.strptime(ts, "%Y%m%d_%H%M%S").isoformat()
            except Exception:
                return None

        items = []
        for jp in json_paths:
            base = jp[:-5]  # .json 제거
            xp = f"{base}.xlsx"
            items.append({
                "json": f"{media_url}refine/{os.path.basename(jp)}",
                "xlsx": f"{media_url}refine/{os.path.basename(xp)}" if xp in xlsx_set else None,
                "created_at": _ts_to_iso(base),
            })

        return Response({
            "items": items,
            "latest": items[0] if items else None
        }, status=200)


# 리액트에 초안들 보여주는 기능
# views.py (마지막에 추가)
class RequirementDraftListView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, project_id):
        # 현재 로그인한 사용자의 프로젝트만 조회
        try:
            project = Project.objects.get(project_id=project_id, user=request.user)
        except Project.DoesNotExist:
            return Response({"error": "해당 프로젝트가 존재하지 않거나 권한이 없습니다."},
                            status=status.HTTP_404_NOT_FOUND)

        drafts = RequirementDraft.objects.filter(project=project)

        # React에서 사용하는 형식에 맞춰 변환
        result = []
        for d in drafts:
            result.append({
                "draft_id": d.RequirementDraft_id,
                "type": "Gemini1" if d.generated_by == "gemini_1" else "Gemini2",
                "feature_name": d.feature_name,
                "summary": d.summary,
                "score_by_model": d.score_by_model,
                "content": d.content,  # JSON 문자열
                "created_at": d.created_at,
            })

        return Response(result, status=status.HTTP_200_OK)



# --- Gemini3: 후보 수집 → Gemini 재랭킹 → 보고서 생성(세 모듈 모두 사용) ---
import os, json, time, re
from django.conf import settings
from django.shortcuts import get_object_or_404
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response

from .models import Project, Requirement, SimilarProject, RequirementDraft

# helpers (토큰은 각 모듈 내부에서 .env 로딩/설정)
from .github_crawler import (
    get_github_instance,
    search_repositories,
    get_readme_content,
    matched_keywords_list,
)
from .github_num import make_similarity_prompt, gemini_similarity_eval
from .similarity_analyzer import analyze_similarity as run_similarity_report


class Gemini3RecommendView(APIView):
    """
    POST /api/project/<project_id>/similar-projects/
    body (옵션):
      {
        "override_keywords": ["custom","keyword","list"],
        "top_k": 3,          # 응답/DB 저장 개수(기본 3)
        "eval_limit": 8,     # 🔸 평가 전 컷오프: Gemini 호출 개수 상한
        "sleep": 0.4         # 🔸 Gemini 호출 사이 간격(초)
      }
    동작:
      1) 확정된 Requirement 수집 → 키워드 구성
      2) GitHub 검색/README 수집 (github_crawler.*)
      3) 🔸 (수정) 평가 전 컷오프 후 Gemini 재랭킹 (github_num.*)
      4) SimilarProject DB 저장
      5) github_repositories.json 저장(컷오프 이후 후보들)
      6) features_*.json 생성(최신 gemini_2 정제본 우선, 없으면 폴백 생성)
      7) TF-IDF+Gemini 보고서 생성(similarity_analyzer)
      8) 보고서 파일을 MEDIA_URL로 노출
    """
    permission_classes = [IsAuthenticated]

    # ── helpers ────────────────────────────────────────────────────────────────
    def _one_line(self, s: str) -> str:
        return re.sub(r"\s+", " ", s or "").strip()

    def _extract_basic_keywords(self, requirements, topk: int = 10):
        """feature_name/summary에서 간단 키워드 추출(영문/숫자/하이픈, 3자 이상, 중복제거)"""
        bag = []
        for r in requirements:
            bag.append(self._one_line(getattr(r, "feature_name", "") or ""))
            bag.append(self._one_line(getattr(r, "summary", "") or ""))
        text = " ".join(bag).lower()
        words = re.findall(r"[a-z0-9][a-z0-9\-]{2,}", text)
        stop = {"and","the","for","with","from","that","this","into","your","you",
                "are","can","have","has","how","use","using","based"}
        uniq, seen = [], set()
        for w in words:
            if w in stop:
                continue
            if w not in seen:
                seen.add(w)
                uniq.append(w)
            if len(uniq) >= topk:
                break
        return uniq if uniq else ["ai", "automation"]

    # ── main ──────────────────────────────────────────────────────────────────
    def post(self, request, project_id):
        # 프로젝트 & 확정 요구사항 확인
        project = get_object_or_404(Project, pk=project_id, user=request.user)
        reqs = Requirement.objects.filter(project=project, confirmed_by_user=True)
        if not reqs.exists():
            return Response({"error": "확정된 Requirement가 없습니다."}, status=400)

        # 1) 키워드: override > 간이 추출
        override = request.data.get("override_keywords")
        if isinstance(override, list) and override:
            keywords = [str(x).strip() for x in override if str(x).strip()]
        else:
            keywords = self._extract_basic_keywords(reqs, topk=10)

        # 2) GitHub 후보 수집
        try:
            g = get_github_instance()  # 토큰은 모듈 내부 처리
            repos = search_repositories(g, keywords, max_repos_per_query=5)
        except Exception as e:
            return Response({"error": f"깃허브 검색 실패: {e}"}, status=500)
        if not repos:
            return Response({"error": "깃허브 후보를 찾지 못했습니다."}, status=404)

        # 3) README 로드 + 키워드 2개 이상 매칭 필터
        candidates = []
        for repo in repos:
            try:
                desc = repo.description or ""
                readme = get_readme_content(repo) or ""
                combined = f"{desc}\n{readme}"
                matched = matched_keywords_list(combined, keywords)
                if len(matched) >= 2:
                    candidates.append({
                        "name": repo.full_name,
                        "url": repo.html_url,
                        "language": repo.language or "",
                        "stars": int(getattr(repo, "stargazers_count", 0) or 0),
                        "description": desc,
                        "readme": readme,
                        "matched_keywords": matched,
                        "matched_count": len(matched),
                    })
            except Exception:
                continue
        if not candidates:
            return Response({"error": "키워드 2개 이상 매칭되는 후보가 없습니다."}, status=404)

        # 🔸 3-b) 평가 전 컷오프: 매칭수/스타 기준 상위 eval_limit 개만 Gemini 호출
        candidates.sort(key=lambda x: (-x.get("matched_count", 0), -x.get("stars", 0)))
        eval_limit = int(request.data.get("eval_limit", 8))
        candidates = candidates[:eval_limit]

        # 4) Gemini 점수화 + 코멘트 (github_num.*)
        merged = "\n".join(f"- {r.feature_name} — {r.summary or ''}" for r in reqs)[:8000]
        sleep_sec = float(request.data.get("sleep", 0.4))
        for c in candidates:
            desc = c.get("description", "") or ""
            readme_trunc = (c.get("readme", "") or "")[:3500]
            prompt = make_similarity_prompt(merged, desc, readme_trunc)
            try:
                score, comment = gemini_similarity_eval(prompt)  # (float|None, str)
            except Exception as e:
                msg = str(e).lower()
                # 친절한 힌트(429/401) + 휴리스틱 폴백
                if "429" in msg or "rate" in msg or "quota" in msg or "exceeded" in msg:
                    time.sleep(3)
                    try:
                        score, comment = gemini_similarity_eval(prompt)
                    except Exception:
                        score = 0.5 * c["matched_count"] + (c["stars"] / 1000.0)
                        comment = "429 지속 → 휴리스틱 점수 사용"
                elif "401" in msg or "unauthorized" in msg or "invalid" in msg or "permission" in msg:
                    score, comment = 0.0, "Gemini 401(키 인증 실패). GEMINI_API_KEY_3 확인."
                else:
                    score, comment = 0.0, f"Gemini 평가 실패: {e}"
            c["gemini_score"] = float(score if score is not None else 0.0)
            c["gemini_comment"] = comment
            time.sleep(sleep_sec)  # 쿼터 보호

        # 5) 정렬 & 저장
        top_k = int(request.data.get("top_k", 3))
        candidates.sort(key=lambda x: (-x.get("gemini_score", 0.0), -x.get("stars", 0)))
        topN = candidates[:top_k]

        SimilarProject.objects.filter(project=project).delete()
        items = []
        for rank, item in enumerate(topN, start=1):
            sp = SimilarProject.objects.create(
                project=project,
                repo_name=item["name"],
                repo_url=item["url"],
                language=item.get("language", "")[:50],
                stars=item.get("stars", 0),
                similarity_score=item.get("gemini_score", 0.0),
            )
            items.append({
                "id": sp.pk,
                "repo_name": sp.repo_name,
                "repo_url": sp.repo_url,
                "language": sp.language,
                "stars": sp.stars,
                "similarity_score": sp.similarity_score,
                "rank": rank,
                "gemini_comment": item.get("gemini_comment", ""),
            })

        # 6) 보고서 입력 파일 저장(분석기는 여기서 로드)
        try:
            with open("github_repositories.json", "w", encoding="utf-8") as f:
                json.dump(candidates, f, ensure_ascii=False, indent=2)

            refined = (
                RequirementDraft.objects
                .filter(project=project, source="gemini_2")
                .order_by("-created_at", "-RequirementDraft_id")
                .first()
            )
            if refined:
                try:
                    refined_json = json.loads(refined.content)
                except Exception:
                    refined_json = refined.content  # 문자열이면 그대로
                payload = {"정제기획서": refined_json}
            else:
                feats = []
                for r in reqs:
                    feats.append({
                        "기능ID": "",
                        "기능명": r.feature_name,
                        "기능설명": {"목적": r.summary, "핵심역할": ""},
                        "사용자시나리오": {"상황": "", "행동": ""},
                        "입력값": {"필수": [], "선택": [], "형식": ""},
                        "출력값": {"요약정보": r.summary, "상세정보": ""},
                        "처리방식": {"단계": [], "사용모델": ""},
                        "예외조건및처리": {"입력누락": "", "오류": ""},
                        "의존성또는연동항목": [],
                        "기능우선순위": "",
                        "UI요소": [],
                        "테스트케이스예시": []
                    })
                payload = {"정제기획서": feats}

            from glob import glob
            nums = [int(p.split("_")[-1].split(".")[0]) for p in glob("features_*.json")
                    if p.split("_")[-1].split(".")[0].isdigit()]
            next_idx = (max(nums) + 1) if nums else 1
            features_path = f"features_{next_idx}.json"
            with open(features_path, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False, indent=2)
        except Exception as e:
            return Response({"error": f"보고서 입력 파일 저장 실패: {e}"}, status=500)

        # 7) 유사도 보고서 생성(similarity_analyzer)
        report_fs_path = None
        try:
            run_similarity_report()  # analysis_report.md 생성
            report_fs_path = os.path.abspath("analysis_report.md")
        except Exception:
            report_fs_path = None  # 보고서 실패해도 추천 결과는 반환

        # 8) MEDIA로 복사해 URL 제공
        report_url = None
        if report_fs_path and os.path.exists(report_fs_path):
            try:
                media_root = getattr(settings, "MEDIA_ROOT",
                                     os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media"))
                media_url = getattr(settings, "MEDIA_URL", "/media/")
                out_dir = os.path.join(media_root, "reports")
                os.makedirs(out_dir, exist_ok=True)

                base = f"project{project.project_id}_analysis_report.md"
                dst = os.path.join(out_dir, base)
                with open(report_fs_path, "rb") as src, open(dst, "wb") as dstf:
                    dstf.write(src.read())
                report_url = f"{media_url}reports/{base}"
            except Exception:
                report_url = None

        return Response({
            "message": "Gemini 3 유사 프로젝트 추천 + 보고서 생성 완료",
            "keywords": keywords,
            "items": items,
            "report_url": report_url
        }, status=201)

    
# (views.py) SimilarProjectListView 교체
class SimilarProjectListView(APIView):
    """
    GET /api/project/<project_id>/similar-projects/list/
    저장된 추천 목록 반환(점수/별점 순)
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, project_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)
        items = SimilarProject.objects.filter(project=project).order_by("-similarity_score", "-stars")

        def _label(s: float) -> str:
            if s is None: return "N/A"
            if s >= 0.85: return "매우 높음"
            if s >= 0.70: return "높음"
            if s >= 0.50: return "보통"
            if s >  0.00: return "낮음"
            return "0점(검토필요)"

        data = []
        for sp in items:
            s = float(getattr(sp, "similarity_score", 0.0) or 0.0)
            pct = round(max(0.0, min(1.0, s)) * 100.0, 1)  # 0~100%
            data.append({
                "id": sp.pk,
                "repo_name": sp.repo_name,
                "repo_url": sp.repo_url,
                "language": sp.language,
                "stars": sp.stars,
                "similarity_score": s,           # 원본 점수
                "score_pct": pct,                # 0~100, 소수1자리
                "score_str": f"{pct}%",          # 표시용 문자열
                "score_label": _label(s),        # 라벨
            })
        return Response(data, status=200)



# views.py (하단에 추가)
from .models import TeamMember, TaskAssignment, Requirement
from rest_framework.permissions import IsAuthenticated
from django.shortcuts import get_object_or_404

# 1. 팀원 추가 (email optional)
from rest_framework.parsers import JSONParser, FormParser, MultiPartParser

class TeamMemberCreateView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [JSONParser, FormParser, MultiPartParser]

    def post(self, request, project_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)

        name   = (request.data.get('name') or '').strip()
        role   = (request.data.get('role') or '').strip()
        skills = (request.data.get('skills') or '').strip()
        email  = (request.data.get('email') or '').strip()   # ← 선택값

        # email은 제외하고 필수 검증
        if not all([name, role, skills]):
            return Response({"error": "이름, 포지션, 기술스택은 필수입니다."}, status=400)

        member = TeamMember.objects.create(
            project=project,
            name=name,
            role=role,
            skills=skills,
            email=email  # 비어 있어도 저장
        )

        return Response(
            {
                "message": "팀원 추가 완료",
                "member_id": member.pk,
                "name": member.name,
                "role": member.role,
                "skills": member.skills,
                "email": member.email,  # 빈 문자열일 수 있음
            },
            status=201
        )



# 2. 팀원 목록 조회
class TeamMemberListView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, project_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)
        members = TeamMember.objects.filter(project=project)
        data = [{
            "id": m.TeamMember,
            "name": m.name,
            "role": m.role,
            "skills": m.skills,
            "email": m.email
        } for m in members]
        return Response(data)


# 3. 팀원 수정
class TeamMemberUpdateView(APIView):
    permission_classes = [IsAuthenticated]

    def put(self, request, project_id, member_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)
        member = get_object_or_404(TeamMember, pk=member_id, project=project)

        member.name = request.data.get("name", member.name)
        member.role = request.data.get("role", member.role)
        member.skills = request.data.get("skills", member.skills)
        member.email = request.data.get("email", member.email)
        member.save()

        return Response({"message": "팀원 정보 수정 완료"})


# 4. 역할 자동 분배 (TaskAssignment 생성)
# views.py 내 기존 AutoAssignTasksView 를 아래 코드로 교체

from rest_framework.parsers import JSONParser, FormParser, MultiPartParser
from django.db import transaction

class AutoAssignTasksView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [JSONParser, FormParser, MultiPartParser]  # JSON/폼/멀티파트 모두 허용

    def post(self, request, project_id):
        import re, json
        from collections import defaultdict

        project = get_object_or_404(Project, pk=project_id, user=request.user)
        members = list(TeamMember.objects.filter(project=project))
        # 확정된 요건만 배정 대상으로 사용
        requirements = list(
            Requirement.objects.filter(project=project, confirmed_by_user=True).order_by("Requirement")
        )

        if not members:
            return Response({"error": "팀원이 없습니다."}, status=400)
        if not requirements:
            return Response({"error": "확정된 기능 명세서가 없습니다."}, status=400)

        # 기존 배정 삭제 (keep=true면 유지)
        keep = str(request.query_params.get("keep", "false")).lower() in ("1", "true", "yes")
        if not keep:
            TaskAssignment.objects.filter(member__project=project).delete()

        # ──────────────────────────────────────────────
        # 정규화 / 토큰화 유틸
        # ──────────────────────────────────────────────
        SYNONYM_MAP = {
            # 언어/런타임
            r"\bc\+\+\b": "cpp", r"\bc sharp\b|\bc#\b": "csharp", r"\bpy(thon)?\b": "python",
            r"\bjs\b|\bjavascript\b": "javascript", r"\bts\b|\btypescript\b": "typescript",
            r"\bnode\.?js\b|\bnodejs\b|\bnode\b": "node", r"\bjava\b": "java",
            r"\bgo(lang)?\b": "golang", r"\brust\b": "rust", r"\bphp\b": "php",
            r"\bkotlin\b": "kotlin", r"\bswift\b": "swift",
            # 프레임워크 / FE
            r"\bdjango\b": "django", r"\bfastapi\b": "fastapi", r"\bspring\b|\bspringboot\b": "spring",
            r"\breact ?native\b": "reactnative", r"\breact\b": "react", r"\bvue(js)?\b": "vue",
            r"\bangular\b": "angular", r"\bnext\.?js\b": "nextjs", r"\btailwind\b": "tailwind",
            # 데이터/DB
            r"\bmysql\b": "mysql", r"\bpostgres(ql)?\b": "postgres", r"\bmaria(db)?\b": "mariadb",
            r"\bredis\b": "redis", r"\bkafka\b": "kafka", r"\bsql\b": "sql", r"\brdbms\b": "db",
            r"\bmongo(db)?\b": "mongodb",
            # 인프라/클라우드/DevOps
            r"\baws\b": "aws", r"\bgcp\b": "gcp", r"\bazure\b": "azure",
            r"\bk8s\b|\bkubernetes\b": "kubernetes", r"\bdocker\b": "docker",
            r"\bterraform\b": "terraform", r"\bci/?cd\b|\bpipeline\b": "cicd",
            # AI/데이터
            r"\bai\b|\bml\b|\bmachine learning\b": "ml",
            r"\bpytorch\b": "pytorch", r"\btensorflow\b|\btf\b": "tensorflow",
            r"\bnlp\b": "nlp", r"\bcv\b|\bcomputer vision\b": "cv",
            r"\bllm\b": "llm", r"\binference\b": "inference", r"\bembedding(s)?\b": "embedding",
            # 기타 도메인/역할
            r"\bbackend\b|백엔드": "backend", r"\bfrontend\b|프론트(엔드)?": "frontend",
            r"\bfull[- ]?stack\b": "fullstack",
            r"\bdevops\b|인프라": "devops", r"\bqa\b|\btest(ing)?\b|테스트": "qa",
            r"\bpm\b|\bproduct manager\b|기획|문서|문서화|스펙|요구사항": "docs",
            r"\bmobile\b|안드로이드|iOS": "mobile",
            r"\bui\b|\bux\b|디자인": "design",
            r"\bapi\b|rest|grpc|msa|microservice": "api",
            r"데이터|분석|통계|bi|warehouse|etl|spark|airflow": "data"
        }

        CATEGORY_WORDS = {
            "backend": {"backend", "api", "server", "db", "sql", "django", "spring", "node", "fastapi", "redis", "kafka"},
            "frontend": {"frontend", "ui", "ux", "react", "vue", "angular", "nextjs", "tailwind"},
            "ai": {"ml", "ai", "llm", "pytorch", "tensorflow", "nlp", "cv", "inference", "embedding"},
            "devops": {"devops", "aws", "gcp", "azure", "kubernetes", "docker", "cicd", "terraform"},
            "data": {"data", "etl", "spark", "airflow", "warehouse", "bi", "db", "sql"},
            "mobile": {"mobile", "android", "ios", "kotlin", "swift", "reactnative", "flutter"},
            "qa": {"qa", "test", "testing", "pytest", "selenium"},
            "docs": {"docs", "문서", "문서화", "스펙", "요구사항", "시나리오"},
            "design": {"design", "ui", "ux"}
        }

        def _norm_text(s: str) -> str:
            return (s or "").lower()

        def _to_tokens(text: str) -> set:
            t = _norm_text(text)
            # 동의어 정규화
            for pattern, rep in SYNONYM_MAP.items():
                t = re.sub(pattern, rep, t, flags=re.IGNORECASE)
            # 토큰 분리(한글/영문/숫자/+# 유지)
            toks = set([tok for tok in re.split(r"[^a-z0-9가-힣\+#]+", t) if len(tok) >= 2])
            return toks

        def _member_profile(m: TeamMember) -> dict:
            # skills는 "Python,Django,MySQL" 혼용 허용
            skills_text = f"{m.skills or ''} {m.role or ''}"
            tokens = _to_tokens(skills_text)
            return {
                "obj": m,
                "tokens": tokens,
                "role": _norm_text(m.role or ""),
                "id": getattr(m, "TeamMember", getattr(m, "pk", None))
            }

        def _flatten_strings(x, bag):
            if isinstance(x, dict):
                for v in x.values(): _flatten_strings(v, bag)
            elif isinstance(x, list):
                for v in x: _flatten_strings(v, bag)
            elif isinstance(x, str):
                bag.append(x)

        def _req_tokens_and_category(req: Requirement):
            pieces = [req.feature_name or "", req.summary or ""]
            # description에 원본 JSON이 들어있으면 펼쳐서 문자열만 수집
            try:
                raw = json.loads(req.description or "{}")
                _flatten_strings(raw, pieces)
            except Exception:
                pass
            toks = _to_tokens(" ".join(pieces))

            # 카테고리 추정
            category = None
            for cat, words in CATEGORY_WORDS.items():
                if toks & words:
                    category = cat
                    break
            return toks, category

        profiles = [_member_profile(m) for m in members]
        if not profiles:
            return Response({"error": "프로필을 만들 수 없습니다."}, status=500)

        # ─ 멤버 기본 정보(그룹 응답용)
        member_info = {
            pf["id"]: {
                "name": pf["obj"].name,
                "position": pf["obj"].role,
                "skills": pf["obj"].skills,
                "email": pf["obj"].email,
            } for pf in profiles
        }

        # 작업량(공평 배분 가중치)
        load = defaultdict(int)
        # 라운드로빈 폴백 포인터
        rr_idx = 0

        created_rows = []
        results = []
        grouped = defaultdict(list)  # member_id -> [tasks...]

        # 트랜잭션: 전부 성공/전부 롤백
        with transaction.atomic():
            for req in requirements:
                req_tokens, req_cat = _req_tokens_and_category(req)

                best = None
                best_score = -9999.0

                for pf in profiles:
                    overlap = len(req_tokens & pf["tokens"])                 # 핵심: 기술 교집합 수
                    role_bonus = 0

                    # 카테고리-역할 이름 가점
                    role_name = pf["role"]
                    if req_cat == "backend" and ("backend" in role_name or "백엔드" in role_name):
                        role_bonus += 2
                    if req_cat == "frontend" and ("frontend" in role_name or "프론트" in role_name):
                        role_bonus += 2
                    if req_cat == "ai" and any(k in role_name for k in ("ai", "ml", "데이터")):
                        role_bonus += 2
                    if req_cat == "devops" and any(k in role_name for k in ("devops", "infra", "ops")):
                        role_bonus += 2
                    if req_cat == "data" and "data" in role_name:
                        role_bonus += 2
                    if req_cat == "mobile" and any(k in role_name for k in ("mobile", "android", "ios")):
                        role_bonus += 2
                    if req_cat == "docs" and any(k in role_name for k in ("pm", "기획", "문서")):
                        role_bonus += 1

                    # 카테고리 키워드가 스킬 토큰에 직접 있으면 추가 가점
                    if req_cat and (CATEGORY_WORDS[req_cat] & pf["tokens"]):
                        role_bonus += 1

                    # 공평성: 이미 많이 배정된 사람은 소폭 감점
                    fairness_penalty = 0.1 * load[pf["id"]]

                    score = overlap * 2 + role_bonus - fairness_penalty
                    if score > best_score:
                        best_score = score
                        best = pf

                # 완전 불일치 → 라운드로빈 폴백
                if best is None or best_score <= 0:
                    best = profiles[rr_idx % len(profiles)]
                    rr_idx += 1

                member = best["obj"]
                load[best["id"]] += 1

                assignment = TaskAssignment.objects.create(
                    requirement=req,
                    member=member,
                    auto_assigned=True
                )
                assignment_id = getattr(assignment, "TaskAssignment", getattr(assignment, "pk", None))
                requirement_id = getattr(req, "Requirement", getattr(req, "pk", None))
                member_id = getattr(member, "TeamMember", getattr(member, "pk", None))

                # 표시용 역할 텍스트(= 과업 타이틀)
                role_text = (req.feature_name or req.summary or "").strip() or "할 일"

                created_rows.append({
                    "assignment_id": assignment_id,
                    "requirement_id": requirement_id,
                    "member_id": member_id
                })

                # 행 단위(기존)
                item = {
                    "assignment_id": assignment_id,
                    "requirement_id": requirement_id,
                    "member_id": member_id,
                    "requirement": req.feature_name,   # 과업 제목
                    "role": role_text,                 # 👈 역할 입력칸 기본값
                    "category": req_cat,
                    "assigned_to": member.name,
                    "score": round(best_score, 2),
                    "matched_skills": sorted(list((_to_tokens(member.skills or "") & req_tokens)))[:10]
                }
                results.append(item)

                # 멤버별 그룹(프론트가 원하는 구조)
                grouped[member_id].append({
                    "assignment_id": assignment_id,
                    "requirement_id": requirement_id,
                    "title": req.feature_name,   # 과업 제목
                    "role": role_text,           # 역할 텍스트
                    "category": req_cat,
                    "score": round(best_score, 2),
                    "matched_skills": item["matched_skills"]
                })

        # 요약(프론트 토스트/배지용)
        by_member = defaultdict(int)
        for r in created_rows:
            by_member[r["member_id"]] += 1
        summary = [{"member_id": mid, "count": c} for mid, c in by_member.items()]

        # 멤버별 그룹 응답 변환
        assignments_grouped = []
        for mid, tasks in grouped.items():
            info = member_info.get(mid, {})
            assignments_grouped.append({
                "member_id": mid,
                "name": info.get("name"),
                "position": info.get("position"),
                "skills": info.get("skills"),
                "email": info.get("email"),
                "task_count": len(tasks),
                "tasks": tasks
            })

        return Response({
            "message": "역할 자동 배정 완료(스킬/스택 기반)",
            "keep_previous": keep,
            "created_count": len(created_rows),
            "created": created_rows,            # id-only 목록
            "summary_by_member": summary,       # 멤버별 생성 개수
            "assignments": results,             # 행 단위(기존 호환)
            "assignments_grouped": assignments_grouped  # 👈 멤버별 묶음(프론트 표시용)
        }, status=201)




# 프로젝트 삭제
class ProjectDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, project_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)
        project.delete()
        return Response({"message": "프로젝트가 성공적으로 삭제되었습니다."}, status=status.HTTP_204_NO_CONTENT)


# 팀원 삭제
class TeamMemberDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, project_id, member_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)
        member = get_object_or_404(TeamMember, pk=member_id, project=project)
        member.delete()
        return Response({"message": "팀원이 삭제되었습니다."}, status=status.HTTP_204_NO_CONTENT)




# ====== 간트차트 생성 ======
# --- 필요한 import (중복되면 하나만 두세요) ---
import os
import json
from datetime import datetime

from django.conf import settings
from django.shortcuts import get_object_or_404
from rest_framework import permissions
from rest_framework.views import APIView
from rest_framework.response import Response

from .models import Project, Requirement, GanttChart, GanttTask, OutputDocument
from .gemini_gantt import make_prompt, call_gemini, parse_llm_array, build_gantt_xlsx


# =========================
# 보조 유틸
# =========================
def _media_subdir(sub: str = "gantt") -> str:
    """MEDIA_ROOT/<sub> 경로 보장 후 절대경로 반환."""
    base = getattr(settings, "MEDIA_ROOT", os.path.join(settings.BASE_DIR, "media"))
    outdir = os.path.join(base, sub)
    os.makedirs(outdir, exist_ok=True)
    return outdir


def _unique_name(outdir: str, prefix: str = "간트차트", ext: str = ".xlsx") -> str:
    """동일 파일명 있으면 _1, _2…로 뒤에 번호를 붙여 유니크한 이름 생성."""
    name = f"{prefix}{ext}"
    i = 1
    while os.path.exists(os.path.join(outdir, name)):
        name = f"{prefix}_{i}{ext}"
        i += 1
    return name


def _sanitize_filename(name: str) -> str:
    """파일명에 쓸 수 없는 문자를 정리."""
    name = (name or "").strip()
    for ch in r'\/:*?"<>|':
        name = name.replace(ch, "_")
    return name[:80] or "간트차트"


def _requirements_payload(project, reqs):
    """간트 프롬프트로 보낼 간단한 페이로드 구성."""
    items = []
    for r in reqs:
        try:
            src = json.loads(r.description) if r.description else {}
        except Exception:
            src = {}
        items.append({
            "기능ID": getattr(r, "Requirement", None),
            "기능명": r.feature_name,
            "요약": r.summary,
            "원본": src,
        })
    return {
        "project": {"title": project.title, "description": project.description},
        "features": items,
    }


# =========================
# imports
# =========================
import os
import re
from datetime import datetime

from django.conf import settings
from django.shortcuts import get_object_or_404

from rest_framework import permissions
from rest_framework.parsers import JSONParser, FormParser, MultiPartParser
from rest_framework.response import Response
from rest_framework.views import APIView

from .models import Project, Requirement, GanttChart, GanttTask

# gemini_gantt.py 유틸 사용
from .gemini_gantt import (
    build_payload_from_db,
    make_prompt,
    call_gemini,
    parse_llm_array,
    build_gantt_xlsx,
    unique_filename,
)

# =========================
# helpers
# =========================
def _sanitize_filename(name: str) -> str:
    """파일명에 부적합 문자를 안전하게 치환"""
    name = (name or "").strip()
    if not name:
        return "간트차트"
    # 한글/영문/숫자/공백/대시/언더스코어만 허용
    name = re.sub(r"[^\w\u3131-\u318E\uAC00-\uD7A3\s\-_]+", "_", name)
    return name[:120]

def _media_subdir(subdir: str) -> str:
    """MEDIA_ROOT/<subdir> 생성 후 절대경로 반환"""
    media_root = getattr(
        settings,
        "MEDIA_ROOT",
        os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media"),
    )
    outdir = os.path.join(media_root, subdir)
    os.makedirs(outdir, exist_ok=True)
    return outdir

def _unique_name(outdir: str, base: str, ext: str) -> str:
    """중복 시 _1, _2… 붙여 유니크 파일명 생성 (gemini_gantt.unique_filename 사용)"""
    return unique_filename(outdir, base=base, ext=ext)

# =========================
# Gantt 생성 뷰 (OutputDocument 미사용) - PATCHED
# =========================
class GanttChartGenerateView(APIView):
    """
    POST /api/project/<int:project_id>/gantt/
    Body(JSON):
    {
      "start_date": "2025-08-12",
      "total_weeks": 12,
      "parts": ["백엔드","프론트엔드","인공지능","서류"],
      "filename": "오토플랜_1차간트"   # optional
    }
    """
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [JSONParser, FormParser, MultiPartParser]

    def post(self, request, project_id):
        import re  # <-- 내부 임포트(외부 수정 불필요)
        project = get_object_or_404(Project, project_id=project_id, user=request.user)

        # 1) 입력값
        start_date_raw = request.data.get("start_date")
        total_weeks = request.data.get("total_weeks")
        parts = request.data.get("parts")
        filename_prefix = _sanitize_filename(request.data.get("filename") or "간트차트")

        # 2) 날짜 파싱 → date
        try:
            sd = str(start_date_raw).strip().replace("/", "-").replace(".", "-")
            start_date_dt = datetime.strptime(sd, "%Y-%m-%d").date()
        except Exception:
            return Response({"error": "start_date 형식 오류 (예: 2025-08-12)"}, status=400)

        # 3) 주차/파트 유효성
        try:
            total_weeks = int(total_weeks)
            if total_weeks <= 0:
                raise ValueError
        except Exception:
            return Response({"error": "total_weeks는 1 이상의 정수여야 합니다."}, status=400)

        if isinstance(parts, str):
            parts = [p.strip() for p in parts.split(",") if p.strip()]
        if not isinstance(parts, list) or not parts:
            return Response({"error": "parts는 비어있지 않은 문자열 배열이어야 합니다."}, status=400)
        parts = [str(p).strip() for p in parts if str(p).strip()]

        # 4) 확정된 Requirement 필수
        reqs = Requirement.objects.filter(project=project, confirmed_by_user=True)
        if not reqs.exists():
            return Response({"error": "확정된 Requirement가 없습니다."}, status=400)

        # 5) LLM 호출/파싱 (gemini_gantt.py)
        payload = build_payload_from_db(project, reqs)
        prompt = make_prompt(payload, parts, total_weeks)
        try:
            llm_text = call_gemini(prompt)
            parsed = parse_llm_array(llm_text)  # JSON 배열
            if not isinstance(parsed, list) or not parsed:
                return Response({"error": "LLM 응답이 유효한 작업 리스트가 아닙니다."}, status=500)
        except Exception as e:
            return Response({"error": f"Gemini 처리 실패: {e}"}, status=500)

        # === 추가: requirement_id 정규화 유틸 ===
        def _to_int_pk(value):
            """문자/라벨('F-001','REQ-12') → 끝자리 숫자 추출해서 int로. 못 찾으면 None."""
            if value is None:
                return None
            if isinstance(value, int):
                return value
            s = str(value).strip()
            if s.isdigit():  # "17"
                return int(s)
            m = re.search(r"(\d+)$", s)  # "F-001" → "001"
            return int(m.group(1)) if m else None

        # === 추가: Requirement dict 캐시(루프마다 쿼리 금지) ===
        req_map = {getattr(r, "Requirement", r.pk): r for r in reqs.only("Requirement")}

        # 6) 간트 헤더 저장
        version = GanttChart.objects.filter(project=project).count() + 1
        gantt = GanttChart.objects.create(
            project=project,
            start_date=start_date_dt,
            total_weeks=total_weeks,
            parts=parts,                   # JSONField라면 그대로 저장
            generated_by="gemini_3",
            source_text=llm_text,
            version=version,
        )

        # 7) 태스크 벌크 저장 (PATCH: requirement 안전 매핑 + parts 안전 처리)
        order_counter, tasks_to_create = {}, []
        for item in parsed:
            # 기능명/작업명 잡기(여러 키 지원)
            feature = (
                item.get("기능명")
                or item.get("feature_name")
                or item.get("작업명")
                or item.get("task")
                or str(item.get("기능ID") or "작업")
            )

            # 파트가 문자열로 올 수도 있음 → 리스트로 정규화
            parts_in_item = item.get("파트") or item.get("part") or item.get("parts") or ["기타"]
            if isinstance(parts_in_item, str):
                parts_in_item = [p.strip() for p in parts_in_item.split(",") if p.strip()]
            if not isinstance(parts_in_item, list) or not parts_in_item:
                parts_in_item = ["기타"]
            parts_in_item = [str(p).strip() for p in parts_in_item if str(p).strip()]

            # 주차/기간 정수화 + 범위 보정
            try:
                start_w = int(item.get("시작주차") or item.get("start_week") or 1)
            except Exception:
                start_w = 1
            try:
                dur_w = int(item.get("기간") or item.get("duration") or item.get("duration_weeks") or 1)
            except Exception:
                dur_w = 1
            if start_w < 1: start_w = 1
            if dur_w < 1: dur_w = 1
            if start_w > total_weeks: start_w = total_weeks
            if start_w + dur_w - 1 > total_weeks:
                dur_w = max(1, total_weeks - start_w + 1)

            # requirement FK 안전 매핑 (핵심 수정)
            rid_raw = item.get("requirement_id") or item.get("기능ID")
            rid_int = _to_int_pk(rid_raw)
            req_obj = req_map.get(rid_int) if rid_int is not None else None

            for p in parts_in_item:
                key = str(p) or "기타"
                order_counter[key] = order_counter.get(key, 0) + 1
                tasks_to_create.append(GanttTask(
                    gantt_chart=gantt,
                    requirement=req_obj,          # None 허용
                    part=key,
                    feature_name=str(feature),
                    start_week=start_w,
                    duration_weeks=dur_w,
                    order=order_counter[key],
                ))

        if tasks_to_create:
            GanttTask.objects.bulk_create(tasks_to_create)

        # 8) 엑셀(.xlsx) 생성  ← 파일명에 projectID 자동 포함
        outdir = _media_subdir("gantt")
        base_prefix = f"project{project.project_id}_{filename_prefix}"
        fname = _unique_name(outdir, base_prefix, ".xlsx")
        abspath = os.path.join(outdir, fname)
        try:
            build_gantt_xlsx(parsed, total_weeks, parts, abspath)
        except Exception as e:
            return Response({"error": f"엑셀 생성 실패: {e}"}, status=500)

        # 파일 경로 저장(상대 경로)
        gantt.file_path = os.path.join("gantt", fname)
        gantt.save(update_fields=["file_path"])

        # ✅ 응답 (G1 스타일: 생성 안내 문구 + files.xlsx)
        media_url = getattr(settings, "MEDIA_URL", "/media/")
        xlsx_url = f"{media_url}{gantt.file_path}"
        return Response({
            "message": f"간트차트 엑셀 파일이 생성되었습니다: {fname}",
            "doc_id": gantt.GanttChart,              # 레거시 호환
            "gantt_id": gantt.GanttChart,            # 권장 필드
            "filename": fname,                        # 예) project6_오토플랜_1차간트.xlsx
            "relative_path": gantt.file_path,         # 예) gantt/project6_...
            "files": {"xlsx": xlsx_url},              # ✅ G1과 유사 구조
            "download_url": f"/api/gantt/download/{gantt.GanttChart}/",
            "download_by_name_url": f"/api/gantt/file/{fname}",
            "public_media_url": xlsx_url,
            "warnings": None
        }, status=201)


# ====== 간트차트 다운로드 뷰들 (GanttChart 기반) ======
import os
from urllib.parse import quote

from django.conf import settings
from django.shortcuts import get_object_or_404
from django.http import FileResponse
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions

from .models import GanttChart

XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

def _media_root():
    return getattr(
        settings,
        "MEDIA_ROOT",
        os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media")
    )

# 1) ID 기반 다운로드: 생성 응답의 download_url 과 매칭
#    GET /api/gantt/download/<int:gantt_id>/
class GanttChartDownloadView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, gantt_id: int):
        gantt = get_object_or_404(GanttChart, GanttChart=gantt_id)

        # 권한 확인
        if gantt.project.user != request.user:
            return Response({"error": "권한 없음"}, status=403)

        rel_path = (gantt.file_path or "").strip()  # 예: "gantt/project6_....xlsx"
        if not rel_path:
            return Response({"error": "파일 경로가 비어 있습니다."}, status=404)

        abs_path = os.path.join(_media_root(), rel_path)
        if not os.path.isfile(abs_path):
            return Response({"error": "파일이 존재하지 않습니다."}, status=404)

        filename = os.path.basename(abs_path)
        resp = FileResponse(
            open(abs_path, "rb"),
            as_attachment=True,
            filename=filename,
            content_type=XLSX_MIME,
        )
        # 한글 파일명 호환
        resp.headers["Content-Disposition"] = f"attachment; filename*=UTF-8''{quote(filename)}"
        return resp


# 2) 파일명 기반 다운로드: 생성 응답의 download_by_name_url 과 매칭
#    GET /api/gantt/file/<path:filename>
#    예) /api/gantt/file/project6_오토플랜_1차간트.xlsx
class GanttChartDownloadByNameView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, filename: str):
        # 안전한 파일명만 허용
        safe_name = os.path.basename(filename)
        rel_path = os.path.join("gantt", safe_name)  # DB에 저장된 형식과 동일

        # 파일 소유권 체크: 내 프로젝트에서 만들어진 간트인지 확인
        rec = GanttChart.objects.filter(file_path=rel_path, project__user=request.user).first()
        if not rec:
            return Response({"error": "파일을 찾을 수 없거나 권한이 없습니다."}, status=404)

        abs_path = os.path.join(_media_root(), rel_path)
        if not os.path.isfile(abs_path):
            return Response({"error": "파일이 존재하지 않습니다."}, status=404)

        resp = FileResponse(
            open(abs_path, "rb"),
            as_attachment=True,
            filename=safe_name,
            content_type=XLSX_MIME,
        )
        resp.headers["Content-Disposition"] = f"attachment; filename*=UTF-8''{quote(safe_name)}"
        return resp




# --- Sidebar: 내 프로젝트 목록 ---
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from django.shortcuts import get_object_or_404

from .models import (
    Project, RequirementDraft, Requirement, SimilarProject,
    TeamMember, ProjectTimeline, OutputDocument, GanttChart
)

class SidebarProjectsView(APIView):
    """
    GET /api/sidebar/projects/
    로그인 사용자가 소유한 프로젝트 목록(사이드바용)
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        qs = (
            Project.objects
            .filter(user=request.user)
            .only("project_id", "title", "created_at")
            .order_by("-project_id")
        )
        items = [
            {
                "project_id": p.project_id,
                "title": p.title,
                "created_at": p.created_at,
            }
            for p in qs
        ]
        return Response({"items": items}, status=200)


class ProjectOverviewView(APIView):
    """
    GET /api/project/<project_id>/overview/
    특정 프로젝트의 핵심 정보(좌측에서 클릭 시 우측 패널에 표시)
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, project_id):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)

        # 가벼운 리스트(최대 N개)만 반환: 클라이언트가 더보기 페이징 요청 가능
        drafts = list(
            RequirementDraft.objects
            .filter(project=project)
            .order_by("-created_at")
            .values("RequirementDraft_id", "source", "feature_name", "summary", "score_by_model", "created_at")[:30]
        )
        requirements = list(
            Requirement.objects
            .filter(project=project)
            .order_by("-created_at")
            .values("Requirement", "feature_name", "summary", "confirmed_by_user", "created_at")[:50]
        )
        similar = list(
            SimilarProject.objects
            .filter(project=project)
            .order_by("-similarity_score", "-stars")
            .values("SimilarProject", "repo_name", "repo_url", "language", "stars", "similarity_score", "created_at")[:10]
        )
        team = list(
            TeamMember.objects
            .filter(project=project)
            .order_by("name")
            .values("TeamMember", "name", "role", "email")[:50]
        )
        outputs = list(
            OutputDocument.objects
            .filter(project=project)
            .order_by("-generated_at")
            .values("OutputDocument", "doc_type", "file_path", "generated_at")[:20]
        )
        gantt = list(
            GanttChart.objects
            .filter(project=project)
            .order_by("-created_at")
            .values("GanttChart", "version", "file_path", "start_date", "total_weeks", "created_at")[:5]
        )
        timeline = (
            ProjectTimeline.objects
            .filter(project=project)
            .values("ProjectTimeline", "start_date", "end_date")
            .first()
        )

        stats = {
            "drafts": RequirementDraft.objects.filter(project=project).count(),
            "requirements": Requirement.objects.filter(project=project).count(),
            "similar_projects": SimilarProject.objects.filter(project=project).count(),
            "team_members": TeamMember.objects.filter(project=project).count(),
            "outputs": OutputDocument.objects.filter(project=project).count(),
        }

        return Response({
            "project": {
                "project_id": project.project_id,
                "title": project.title,
                "description": project.description,
                "created_at": project.created_at,
            },
            "stats": stats,
            "drafts": drafts,
            "requirements": requirements,
            "similar_projects": similar,
            "team": team,
            "outputs": outputs,
            "gantt_charts": gantt,
            "timeline": timeline,
        }, status=200)

# views.py 하단에 추가
from rest_framework.permissions import IsAuthenticated
from django.shortcuts import get_object_or_404
from .models import Requirement, SimilarProject, Project

class ConfirmedRequirementListView(APIView):
    """
    GET /api/project/<project_id>/requirements/confirmed/
    확정된 기능 명세서 목록만 반환
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, project_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)
        items = list(
            Requirement.objects
            .filter(project=project, confirmed_by_user=True)
            .order_by("-created_at")
            .values("Requirement", "feature_name", "summary", "description", "created_at")
        )
        return Response(items, status=200)


class ConfirmedAndSimilarView(APIView):
    """
    GET /api/project/<project_id>/requirements/confirmed-and-similar/
    확정된 기능 명세서와 저장된 유사 프로젝트를 한 번에 반환
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, project_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)

        confirmed = list(
            Requirement.objects
            .filter(project=project, confirmed_by_user=True)
            .order_by("-created_at")
            .values("Requirement", "feature_name", "summary", "description", "created_at")
        )
        similar = list(
            SimilarProject.objects
            .filter(project=project)
            .order_by("-similarity_score", "-stars")
            .values("SimilarProject", "repo_name", "repo_url", "language", "stars", "similarity_score", "created_at")
        )

        return Response({
            "confirmed": confirmed,
            "similar": similar
        }, status=200)
    
# 제미나이 아이디어
# views.py
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from django.conf import settings
from django.shortcuts import get_object_or_404
import os
import json
from datetime import datetime

# ⬇ gemini_idea.py의 함수 재사용
from .idea_to_plan_generator import build_similar_map, generate_markdown, generate_word  # :contentReference[oaicite:1]{index=1}

# ⬇ 정제/확장 로직은 원래 gemini_idea.py에서 IdeaRefiner/IdeaExpander가 담당
from .idea_refiner import IdeaRefiner
from .idea_expander import IdeaExpander

# (옵션) 모델을 써서 확정 Requirement에서 core_features를 구성하는 뷰
from .models import Project, Requirement


from .models import Project, Requirement, RequirementDraft, SimilarProject

class IdeaProcessView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        idea = (request.data.get("idea") or "").strip()
        if not idea:
            return Response({"error": "idea(아이디어) 가 비었습니다."}, status=400)

        # 1) 정제
        refiner = IdeaRefiner()
        refined = refiner.refine(idea)

        # 2) 확장
        expander = IdeaExpander()
        suggestions = expander.expand(refined)

        # 3) 유사맵
        core_features = refined.get("core_features") or []
        similar_map = build_similar_map(core_features)

        # 4) === DB 저장 ===
        project_id = request.data.get("project_id")
        project = None
        if project_id:
            project = get_object_or_404(Project, pk=project_id, user=request.user)

            # RequirementDraft 저장
            for feature in core_features:
                RequirementDraft.objects.create(
                    project=project,
                    source="gemini_1",
                    feature_name=feature,
                    summary="",  # 필요 시 refined에서 채움
                    score_by_model=0.0,  # 필요 시 모델 점수
                    content=json.dumps(refined, ensure_ascii=False),
                    generated_by="IdeaProcessView"
                )

            # SimilarProject 저장
            for feature, repo_list in similar_map.items():
                for repo_info in repo_list:
                    for repo in repo_info.get("repos", []):
                        SimilarProject.objects.create(
                            project=project,
                            repo_name=repo,
                            repo_url=f"https://github.com/{repo}",
                            language="",  # 필요 시 추가
                            stars=0,  # 필요 시 추가
                            similarity_score=0.0
                        )

        out = {
            "refined": refined,
            "suggestions": suggestions,
            "similar_map": similar_map,
        }

        if request.data.get("also_return_markdown") is True:
            md = generate_markdown(refined, suggestions, similar_map)
            out["markdown"] = md

        return Response(out, status=200)


class IdeaExportView(APIView):
    """
    POST /api/idea/export/
    Body:
    {
      "refined": {"goal": "...", "target_user": "...", "core_features": [...]},
      "suggestions": ["...", "..."],
      "similar_map": {...},
      "filename_prefix": "idea_plan"   # 선택, 기본 'idea_plan'
    }
    Response:
    {
      "markdown_file": "idea_plan_3.md",
      "docx_file": "idea_plan_3.docx",
      "saved_dir": "/absolute/or/media/path"
    }
    """
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        refined = request.data.get("refined") or {}
        suggestions = request.data.get("suggestions") or []
        similar_map = request.data.get("similar_map") or {}

        # 최소 필드 체크
        if not refined or "goal" not in refined or "core_features" not in refined:
            return Response({"error": "refined(goal, core_features 포함)가 필요합니다."}, status=400)

        # 저장 디렉토리(프로젝트 MEDIA_ROOT/ideas) 보장
        base = getattr(settings, "MEDIA_ROOT", os.path.join(settings.BASE_DIR, "media"))
        outdir = os.path.join(base, "ideas")
        os.makedirs(outdir, exist_ok=True)

        # 파일명 프리픽스
        prefix = (request.data.get("filename_prefix") or "idea_plan").strip()
        # 자동 인덱스
        existing = [f for f in os.listdir(outdir) if f.startswith(prefix + "_")]
        nums = []
        for f in existing:
            try:
                stem = os.path.splitext(f)[0]
                n = stem.split("_")[-1]
                if n.isdigit():
                    nums.append(int(n))
            except Exception:
                pass
        next_idx = max(nums) + 1 if nums else 1
        md_name = f"{prefix}_{next_idx}.md"
        docx_name = f"{prefix}_{next_idx}.docx"

        # 1) MD 내용 생성 후 저장
        md_text = generate_markdown(refined, suggestions, similar_map)  # :contentReference[oaicite:4]{index=4}
        md_path = os.path.join(outdir, md_name)
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_text)

        # 2) DOCX 생성 후 저장
        docx_path = os.path.join(outdir, docx_name)
        generate_word(refined, suggestions, similar_map, docx_path)     # :contentReference[oaicite:5]{index=5}

        return Response({
            "markdown_file": md_name,
            "docx_file": docx_name,
            "saved_dir": outdir
        }, status=201)


class IdeaFromConfirmedRequirementsView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, project_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)
        reqs = Requirement.objects.filter(project=project, confirmed_by_user=True).order_by("created_at")

        if not reqs.exists():
            return Response({"error": "확정된 Requirement가 없습니다."}, status=404)

        core_features = [r.feature_name for r in reqs]
        refined = {
            "goal": project.title,
            "target_user": "",
            "core_features": core_features,
        }
        suggestions = []
        similar_map = build_similar_map(core_features)

        # === DB 저장 ===
        # SimilarProject 저장
        for feature, repo_list in similar_map.items():
            for repo_info in repo_list:
                for repo in repo_info.get("repos", []):
                    SimilarProject.objects.create(
                        project=project,
                        repo_name=repo,
                        repo_url=f"https://github.com/{repo}",
                        language="",
                        stars=0,
                        similarity_score=0.0
                    )

        payload = {
            "refined": refined,
            "suggestions": suggestions,
            "similar_map": similar_map
        }

        # (옵션) export
        if request.data.get("also_export") is True:
            base = getattr(settings, "MEDIA_ROOT", os.path.join(settings.BASE_DIR, "media"))
            outdir = os.path.join(base, "ideas")
            os.makedirs(outdir, exist_ok=True)
            prefix = (request.data.get("filename_prefix") or "idea_plan").strip()

            existing = [f for f in os.listdir(outdir) if f.startswith(prefix + "_")]
            nums = []
            for f in existing:
                try:
                    stem = os.path.splitext(f)[0]
                    n = stem.split("_")[-1]
                    if n.isdigit():
                        nums.append(int(n))
                except Exception:
                    pass
            next_idx = max(nums) + 1 if nums else 1
            md_name = f"{prefix}_{next_idx}.md"
            docx_name = f"{prefix}_{next_idx}.docx"

            md_text = generate_markdown(refined, suggestions, similar_map)
            with open(os.path.join(outdir, md_name), "w", encoding="utf-8") as f:
                f.write(md_text)
            generate_word(refined, suggestions, similar_map, os.path.join(outdir, docx_name))

            payload["export"] = {
                "markdown_file": md_name,
                "docx_file": docx_name,
                "saved_dir": outdir
            }

        return Response(payload, status=200)

# ===== Sidebar Tree (프로젝트 → 하위 항목들) =====
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from django.shortcuts import get_object_or_404

from .models import (
    Project, RequirementDraft, Requirement, SimilarProject,
    TeamMember, OutputDocument, GanttChart
)
from django.conf import settings
import os

def _gantt_doc_id_by_path(project, path):
    """간트차트 파일경로로 OutputDocument(pk) 찾기 (다운로드 URL 만들 때 사용)"""
    if not path:
        return None
    doc = OutputDocument.objects.filter(project=project, file_path=path).order_by("-OutputDocument").first()
    return getattr(doc, "OutputDocument", None) if doc else None

def _project_tree_dict(project, limit_each=8):
    """프로젝트 1개에 대한 트리 구조 생성"""
    pid = getattr(project, "project_id", project.pk)

    # Drafts
    drafts_qs = (
        RequirementDraft.objects
        .filter(project=project)
        .order_by("-created_at", "-RequirementDraft_id")
        .values("RequirementDraft_id", "source", "feature_name", "summary", "created_at")[:limit_each]
    )
    drafts = [{
        "type": "draft",
        "id": d["RequirementDraft_id"],
        "label": d["feature_name"] or f"Draft {d['RequirementDraft_id']}",
        "summary": d["summary"],
        "source": d["source"],
        "url": f"/api/project/{pid}/drafts/{d['RequirementDraft_id']}/"
    } for d in drafts_qs]
    drafts_count = RequirementDraft.objects.filter(project=project).count()

    # Confirmed Requirements
    reqs_qs = (
        Requirement.objects
        .filter(project=project, confirmed_by_user=True)
        .order_by("-created_at", "-Requirement")
        .values("Requirement", "feature_name", "summary", "created_at")[:limit_each]
    )
    confirmed = [{
        "type": "requirement",
        "id": r["Requirement"],
        "label": r["feature_name"],
        "summary": r["summary"],
        "url": f"/api/requirements/{r['Requirement']}/"
    } for r in reqs_qs]
    confirmed_count = Requirement.objects.filter(project=project, confirmed_by_user=True).count()

    # Similar Projects
    sim_qs = (
        SimilarProject.objects
        .filter(project=project)
        .order_by("-similarity_score", "-stars")
        .values("SimilarProject", "repo_name", "repo_url", "language", "stars")[:limit_each]
    )
    similar = [{
        "type": "similar_project",
        "id": s["SimilarProject"],
        "label": s["repo_name"],
        "language": s["language"],
        "stars": s["stars"],
        "href": s["repo_url"]  # 외부로 바로 이동
    } for s in sim_qs]
    similar_count = SimilarProject.objects.filter(project=project).count()

    # Team Members
    team_qs = (
        TeamMember.objects
        .filter(project=project)
        .order_by("name")
        .values("TeamMember", "name", "role", "email")[:limit_each]
    )
    team = [{
        "type": "member",
        "id": m["TeamMember"],
        "label": m["name"],
        "role": m["role"],
        "email": m["email"],
        "url": f"/api/project/{pid}/team-members/{m['TeamMember']}/"
    } for m in team_qs]
    team_count = TeamMember.objects.filter(project=project).count()

    # Gantt Charts
    gantt_qs = (
        GanttChart.objects
        .filter(project=project)
        .order_by("-created_at", "-GanttChart")
        .values("GanttChart", "version", "file_path", "start_date", "total_weeks")[:limit_each]
    )
    gantts = []
    for g in gantt_qs:
        doc_id = _gantt_doc_id_by_path(project, g["file_path"])
        download_url = f"/api/gantt/download/{doc_id}/" if doc_id else None
        label = f"v{g['version']} · {g['start_date']} 시작 · {g['total_weeks']}주"
        gantts.append({
            "type": "gantt",
            "id": g["GanttChart"],
            "label": label,
            "download_url": download_url,
            "file_path": g["file_path"]
        })
    gantt_count = GanttChart.objects.filter(project=project).count()

    # Outputs (문서 등)
    out_qs = (
        OutputDocument.objects
        .filter(project=project)
        .order_by("-generated_at", "-OutputDocument")
        .values("OutputDocument", "doc_type", "file_path", "generated_at")[:limit_each]
    )
    outputs = [{
        "type": "output",
        "id": o["OutputDocument"],
        "label": o["doc_type"],
        "file_path": o["file_path"],
        "download_hint": o["file_path"]  # 일반 산출물은 별도 다운로드 뷰 없으면 파일 경로로 처리
    } for o in out_qs]
    out_count = OutputDocument.objects.filter(project=project).count()

    return {
        "type": "project",
        "project_id": pid,
        "title": project.title,
        "url": f"/api/project/{pid}/overview/",
        "children": [
            {
                "type": "group", "label": "Drafts", "count": drafts_count,
                "url": f"/api/project/{pid}/drafts/",
                "children": drafts
            },
            {
                "type": "group", "label": "Requirements (Confirmed)", "count": confirmed_count,
                "url": f"/api/project/{pid}/requirements/confirmed/",
                "children": confirmed
            },
            {
                "type": "group", "label": "Gantt Charts", "count": gantt_count,
                "url": f"/api/project/{pid}/gantt/list/",
                "children": gantts
            },
            {
                "type": "group", "label": "Team Members", "count": team_count,
                "url": f"/api/project/{pid}/team-members/",
                "children": team
            },
            {
                "type": "group", "label": "Similar OSS", "count": similar_count,
                "url": f"/api/project/{pid}/similar-projects/list/",
                "children": similar
            },
            {
                "type": "group", "label": "Outputs", "count": out_count,
                "url": f"/api/project/{pid}/outputs/",
                "children": outputs
            }
        ]
    }

class SidebarTreeAllView(APIView):
    """
    GET /api/sidebar/tree/
    로그인 사용자의 모든 프로젝트에 대해, 폴더형 트리 구조 반환
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        projects = Project.objects.filter(user=request.user).order_by("-project_id")
        data = [_project_tree_dict(p, limit_each=8) for p in projects]
        return Response({"projects": data}, status=200)

class SidebarTreeProjectView(APIView):
    """
    GET /api/sidebar/project/<project_id>/tree/
    특정 프로젝트만 트리 구조
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, project_id):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)
        return Response(_project_tree_dict(project, limit_each=20), status=200)

# ====== 세부 항목 조회용 라이트 상세 API들 ======

class RequirementDraftDetailView(APIView):
    """GET /api/project/<project_id>/drafts/<draft_id>/"""
    permission_classes = [IsAuthenticated]
    def get(self, request, project_id, draft_id):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)
        d = get_object_or_404(RequirementDraft, RequirementDraft_id=draft_id, project=project)
        return Response({
            "draft_id": d.RequirementDraft_id,
            "source": d.source,
            "feature_name": d.feature_name,
            "summary": d.summary,
            "score_by_model": d.score_by_model,
            "content": d.content,
            "created_at": d.created_at
        }, status=200)

class RequirementDetailView(APIView):
    """GET /api/requirements/<req_id>/"""
    permission_classes = [IsAuthenticated]
    def get(self, request, req_id):
        r = get_object_or_404(Requirement, Requirement=req_id)
        if r.project.user != request.user:
            return Response({"error": "권한 없음"}, status=403)
        return Response({
            "Requirement": r.Requirement,
            "project_id": getattr(r.project, "project_id", None),
            "feature_name": r.feature_name,
            "summary": r.summary,
            "description": r.description,
            "confirmed_by_user": r.confirmed_by_user,
            "created_at": r.created_at
        }, status=200)

class TeamMemberDetailView(APIView):
    """GET /api/project/<project_id>/team-members/<member_id>/"""
    permission_classes = [IsAuthenticated]
    def get(self, request, project_id, member_id):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)
        m = get_object_or_404(TeamMember, TeamMember=member_id, project=project)
        return Response({
            "id": m.TeamMember,
            "name": m.name,
            "role": m.role,
            "skills": m.skills,
            "email": m.email,
            "project_id": getattr(project, "project_id", None)
        }, status=200)

class GanttChartListView(APIView):
    """GET /api/project/<project_id>/gantt/list/"""
    permission_classes = [IsAuthenticated]
    def get(self, request, project_id):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)
        items = []
        for g in GanttChart.objects.filter(project=project).order_by("-created_at", "-GanttChart"):
            doc_id = _gantt_doc_id_by_path(project, g.file_path)
            items.append({
                "GanttChart": g.GanttChart,
                "version": g.version,
                "file_path": g.file_path,
                "start_date": g.start_date,
                "total_weeks": g.total_weeks,
                "download_url": f"/api/gantt/download/{doc_id}/" if doc_id else None
            })
        return Response(items, status=200)

class OutputDocumentListView(APIView):
    """GET /api/project/<project_id>/outputs/"""
    permission_classes = [IsAuthenticated]
    def get(self, request, project_id):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)
        items = list(
            OutputDocument.objects
            .filter(project=project)
            .order_by("-generated_at", "-OutputDocument")
            .values("OutputDocument", "doc_type", "file_path", "generated_at")
        )
        return Response(items, status=200)
    
# views.py (기존 G1/G2 다운로드 뷰 교체)

from django.conf import settings
from django.http import FileResponse, Http404
from django.shortcuts import get_object_or_404
from rest_framework.views import APIView
from rest_framework import permissions
import os, glob

from .models import Project

EXCEL_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

def _send_xlsx(abspath: str) -> FileResponse:
    if not os.path.exists(abspath):
        raise Http404("엑셀 파일이 존재하지 않습니다.")
    fname = os.path.basename(abspath)
    size = os.path.getsize(abspath)
    f = open(abspath, "rb")
    resp = FileResponse(f, as_attachment=True, filename=fname, content_type=EXCEL_MIME)
    resp["Content-Type"] = EXCEL_MIME
    resp["Content-Disposition"] = f"attachment; filename*=UTF-8''{fname}"
    resp["X-Content-Type-Options"] = "nosniff"
    resp["Cache-Control"] = "no-transform, no-store"
    resp["Content-Transfer-Encoding"] = "binary"
    resp["Accept-Ranges"] = "none"
    resp["Content-Length"] = str(size)
    return resp

def _normalize_ts(ts: str) -> str:
    # 14자리 숫자면 가운데 언더바 삽입
    if ts and ts.isdigit() and len(ts) == 14:
        return ts[:8] + "_" + ts[8:]
    return ts

def _latest_file(path_pattern: str) -> str | None:
    files = sorted(glob.glob(path_pattern), reverse=True)
    return files[0] if files else None

class ProjectG1XlsxDownloadView(APIView):
    """
    GET /api/project/<project_id>/download/g1/<str:ts>/
      - ts: YYYYMMDD_HHMMSS | YYYYMMDDHHMMSS | latest
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, project_id, ts):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)
        media_root = getattr(settings, "MEDIA_ROOT", os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media"))
        drafts_dir = os.path.join(media_root, "drafts")

        if ts == "latest":
            cand = _latest_file(os.path.join(drafts_dir, f"project{project.project_id}_*_g1.xlsx"))
            if not cand:
                raise Http404("G1 엑셀 파일이 없습니다.")
            return _send_xlsx(cand)

        ts = _normalize_ts(ts)
        xlsx_path = os.path.join(drafts_dir, f"project{project.project_id}_{ts}_g1.xlsx")
        if not os.path.exists(xlsx_path):
            # 폴더에서 가장 최신 파일로 한번 더 시도(사용자 편의)
            cand = _latest_file(os.path.join(drafts_dir, f"project{project.project_id}_*_g1.xlsx"))
            if not cand:
                raise Http404("해당 ts의 G1 엑셀 파일이 없습니다.")
            xlsx_path = cand
        return _send_xlsx(xlsx_path)

class ProjectG2XlsxDownloadView(APIView):
    """
    GET /api/project/<project_id>/download/g2/<str:ts>/
      - ts: YYYYMMDD_HHMMSS | YYYYMMDDHHMMSS | latest
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, project_id, ts):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)
        media_root = getattr(settings, "MEDIA_ROOT", os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media"))
        refine_dir = os.path.join(media_root, "refine")

        if ts == "latest":
            cand = _latest_file(os.path.join(refine_dir, f"project{project.project_id}_*_g2_fix.xlsx"))
            if not cand:
                raise Http404("G2 엑셀 파일이 없습니다.")
            return _send_xlsx(cand)

        ts = _normalize_ts(ts)
        xlsx_path = os.path.join(refine_dir, f"project{project.project_id}_{ts}_g2_fix.xlsx")
        if not os.path.exists(xlsx_path):
            cand = _latest_file(os.path.join(refine_dir, f"project{project.project_id}_*_g2_fix.xlsx"))
            if not cand:
                raise Http404("해당 ts의 G2 엑셀 파일이 없습니다.")
            xlsx_path = cand
        return _send_xlsx(xlsx_path)

# --- MD 원문 그대로 반환(미리보기 RAW) ---
import os
from django.http import Http404, HttpResponse
from django.conf import settings
from django.shortcuts import get_object_or_404
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from .models import Project

def _reports_dir():
    media_root = getattr(settings, "MEDIA_ROOT",
                         os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media"))
    path = os.path.join(media_root, "reports")
    os.makedirs(path, exist_ok=True)
    return path

class ProjectReportRawView(APIView):
    """
    GET /api/project/<project_id>/reports/latest/raw/
    -> 최신 보고서를 JSON 래핑 없이 text/markdown으로 바로 반환
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, project_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)
        reports_dir = _reports_dir()
        # 최신 파일 선택 (고정 파일명 사용 중이면 그대로 참조)
        path = os.path.join(reports_dir, f"project{project.project_id}_analysis_report.md")
        if not os.path.isfile(path):
            raise Http404("보고서가 없습니다. 먼저 유사도 분석을 실행하세요.")
        with open(path, "rb") as f:
            data = f.read()
        resp = HttpResponse(data, content_type="text/markdown; charset=utf-8")
        # 첨부로 다운받지 않고 브라우저/프론트에서 바로 렌더링하도록 Content-Disposition 생략
        # CORS가 필요하면 settings의 CORS_ALLOW_HEADERS/ORIGINS 설정을 확인
        resp["Cache-Control"] = "no-store"
        return resp

class ProjectReportRawByNameView(APIView):
    """
    GET /api/project/<project_id>/reports/raw/<str:filename>/
    -> 파일명을 지정해서 원문 반환 (text/markdown)
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, project_id, filename):
        project = get_object_or_404(Project, pk=project_id, user=request.user)
        reports_dir = _reports_dir()
        # 보안: 해당 프로젝트 파일만 허용
        if not filename.startswith(f"project{project.project_id}_") or not filename.endswith(".md"):
            raise Http404("파일 접근 권한이 없거나 잘못된 이름입니다.")
        path = os.path.join(reports_dir, filename)
        if not os.path.isfile(path):
            raise Http404("파일을 찾을 수 없습니다.")
        with open(path, "rb") as f:
            data = f.read()
        resp = HttpResponse(data, content_type="text/markdown; charset=utf-8")
        resp["Cache-Control"] = "no-store"
        return resp


# === 확정 기획서/명세서 기반 파일 생성 뷰 ===
# 기준: Project.description(기획서) + Requirement.confirmed_by_user(True)(명세서) + 최신 G2 초안(있으면)
# 출력: 제안(JSON/MD) 또는 코드 ZIP, 최종 합본 MD

import os, io, json, zipfile
from datetime import datetime
from django.conf import settings
from django.shortcuts import get_object_or_404
from rest_framework import permissions
from rest_framework.views import APIView
from rest_framework.response import Response

from .models import Project, Requirement, RequirementDraft  # 모델 스키마 근거: confirmed_by_user, RequirementDraft_id 등

# ---------- 공통 유틸 ----------
def _media_root():
    return getattr(settings, "MEDIA_ROOT",
                   os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media"))

def _media_url():
    return getattr(settings, "MEDIA_URL", "/media/")

def _ensure_dir(path: str):
    os.makedirs(path, exist_ok=True); return path

def _nowtag():
    return datetime.now().strftime("%Y%m%d_%H%M%S")

def _save_text(rel_dir: str, filename: str, text: str):
    abs_dir = os.path.join(_media_root(), rel_dir); _ensure_dir(abs_dir)
    abs_path = os.path.join(abs_dir, filename)
    with open(abs_path, "w", encoding="utf-8") as f:
        f.write(text or "")
    url = _media_url().rstrip("/") + "/" + os.path.join(rel_dir, filename).replace("\\", "/")
    return abs_path, url

def _save_bytes(rel_dir: str, filename: str, data: bytes):
    abs_dir = os.path.join(_media_root(), rel_dir); _ensure_dir(abs_dir)
    abs_path = os.path.join(abs_dir, filename)
    with open(abs_path, "wb") as f:
        f.write(data or b"")
    url = _media_url().rstrip("/") + "/" + os.path.join(rel_dir, filename).replace("\\", "/")
    return abs_path, url

def _zip_bytes(files):
    bio = io.BytesIO()
    with zipfile.ZipFile(bio, "w", zipfile.ZIP_DEFLATED) as zf:
        for arc_path, content in files:
            zf.writestr(arc_path, content or "")
    bio.seek(0); return bio.read()

# ---------- 확정본 수집 ----------
def _get_confirmed_plan_spec(project: Project):
    """
    기획서: Project.description
    명세서: Requirement.confirmed_by_user=True 만 모아 Markdown 조합
    둘 중 하나라도 없으면 404로 반환하기 위해 ValueError 발생
    """
    plan_md = (project.description or "").strip()  # 기획서(확정): 프로젝트 설명란 사용
    reqs = list(Requirement.objects.filter(project=project, confirmed_by_user=True).order_by("created_at"))
    if not plan_md or not reqs:
        raise ValueError("확정 기획서 또는 확정 명세서가 없습니다.")

    # 확정 명세서 → Markdown 조립
    lines = ["# 기능 명세서(확정)"]
    for r in reqs:
        lines.append(f"\n## {r.feature_name}\n")
        if r.summary:
            lines.append(f"- 요약: {r.summary}")
        if r.description:
            lines.append(f"\n{r.description}\n")
    spec_md = "\n".join(lines).strip()
    return plan_md, spec_md, reqs

def _get_latest_g2_draft(project: Project):
    """
    최신 G2 초안(JSON 문자열 또는 기타 문자열). 없으면 None
    모델 PK명이 RequirementDraft_id 임을 반영해 정렬. 
    """
    d = (RequirementDraft.objects
         .filter(project=project, source="gemini_2")
         .order_by("-created_at", "-RequirementDraft_id")
         .first())
    return d.content if d else None

# ---------- 1) 툴 제안(확정본 기반) ----------
class ProjectToolsProposeFromConfirmedView(APIView):
    """
    POST /api/project/<project_id>/tools/propose/
    body 무시(확정본 기반). 결과:
      media/artifacts/project{pid}/tools_recommendation_*.json
      media/artifacts/project{pid}/tools_summary_*.md
    """
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, project_id):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)
        try:
            plan_md, spec_md, _ = _get_confirmed_plan_spec(project)
        except ValueError as e:
            return Response({"error": str(e)}, status=404)

        # (실제 로직은 tools_generator.propose_tools(plan_md, spec_md)로 분리 가능)
        summary_json = {
            "frontend": {"recommendation": "React.tsx", "reason": "확정 명세서 기반 기본 추천"},
            "backend":  {"recommendation": "Django",     "reason": "ORM/Admin/생태계"},
            "sql":      {"recommendation": "PostgreSQL", "reason": "확장성/무난"},
            "ai":       {"recommendation": "Gemini",     "reason": "기본 모델"},
        }
        md = (
            "## 기술 스택 추천(확정본 기반)\n\n"
            f"- **SQL**: {summary_json['sql']['recommendation']}\n"
            f"- **Backend**: {summary_json['backend']['recommendation']}\n"
            f"- **Frontend**: {summary_json['frontend']['recommendation']}\n"
            f"- **AI(Model)**: {summary_json['ai']['recommendation']}\n"
        )

        rel = f"artifacts/project{project.project_id}"
        ts = _nowtag()
        _, json_url = _save_text(rel, f"tools_recommendation_{ts}.json",
                                 json.dumps(summary_json, ensure_ascii=False, indent=2))
        _, md_url   = _save_text(rel, f"tools_summary_{ts}.md", md)
        return Response({"json_url": json_url, "summary_md_url": md_url}, status=201)

# ---------- 2) 코드/산출물 생성(확정본 기반) ----------
class ProjectGenerateSQLFromConfirmedView(APIView):
    """POST /api/project/<project_id>/generate/sql/ → SQL 스캐폴드 ZIP"""
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, project_id):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)
        try:
            plan_md, spec_md, reqs = _get_confirmed_plan_spec(project)
        except ValueError as e:
            return Response({"error": str(e)}, status=404)

        # (code_generator.generate_sql(...)로 분리 가능)
        files = [
            ("SQL/README.md", "# SQL 산출물(확정본 기반)\n- Project.description + 확정 Requirement 반영\n"),
            ("SQL/schema.sql", "-- TODO: 확정 명세서 기반 DDL\n"),
            ("SQL/seed.sql",   "-- TODO: 샘플 시드 데이터\n"),
        ]
        data = _zip_bytes(files)
        rel = f"artifacts/project{project.project_id}"
        ts = _nowtag()
        _, url = _save_bytes(rel, f"project{project.project_id}_sql_{ts}.zip", data)
        return Response({"zip_url": url}, status=201)

class ProjectGenerateBackendFromConfirmedView(APIView):
    """POST /api/project/<project_id>/generate/backend/ → 백엔드 스캐폴드 ZIP"""
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, project_id):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)
        try:
            plan_md, spec_md, reqs = _get_confirmed_plan_spec(project)
        except ValueError as e:
            return Response({"error": str(e)}, status=404)

        files = [
            ("Back/README.md", "# Backend(확정본 기반)\n- Django 예시\n"),
            ("Back/app/__init__.py", ""),
            ("Back/app/views.py", "from django.http import JsonResponse\n\ndef ping(_):\n    return JsonResponse({'ok': True})\n"),
            ("Back/app/urls.py", "from django.urls import path\nfrom .views import ping\nurlpatterns = [ path('ping/', ping) ]\n"),
            ("Back/manage.py", "# placeholder\n"),
        ]
        data = _zip_bytes(files)
        rel = f"artifacts/project{project.project_id}"
        ts = _nowtag()
        _, url = _save_bytes(rel, f"project{project.project_id}_backend_{ts}.zip", data)
        return Response({"zip_url": url}, status=201)

class ProjectGenerateFrontendFromConfirmedView(APIView):
    """POST /api/project/<project_id>/generate/frontend/ → 프론트 스캐폴드 ZIP"""
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, project_id):
        project = get_object_or_404(Project, project_id=project_id, user=request.user)
        try:
            plan_md, spec_md, reqs = _get_confirmed_plan_spec(project)
        except ValueError as e:
            return Response({"error": str(e)}, status=404)

        files = [
            ("Front/README.md", "# Frontend(확정본 기반)\n- React 예시\n"),
            ("Front/src/App.jsx", "export default function App(){ return (<div>App</div>); }"),
            ("Front/src/main.jsx", "import App from './App.jsx'\nconsole.log('boot')\n"),
            ("Front/package.json", '{"name":"front","private":true}'),
        ]
        data = _zip_bytes(files)
        rel = f"artifacts/project{project.project_id}"
        ts = _nowtag()
        _, url = _save_bytes(rel, f"project{project.project_id}_frontend_{ts}.zip", data)
        return Response({"zip_url": url}, status=201)

# ==== FinalizeRequirementView & FinalDevDocGenerateView (fixed + main.py 통합) ====
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework import status
from django.conf import settings
from django.utils import timezone
from django.shortcuts import get_object_or_404
from django.core.exceptions import FieldError

from pathlib import Path
import os, json, re, tempfile, types

from docx import Document as DocxDocument

from .models import Project, Requirement, RequirementDraft
from .auto_document import (
    load_docx_and_plaintext,
    detect_placeholders_in_text,
    docx_replace_placeholders,
)

# ✅ main.py 파이프라인 병행 사용
try:
    # 프로젝트 루트에 main.py가 있고, 내부에서 src.* 의존성이 충족되어야 합니다.
    from . import main as rd_main  # run_pipeline(args) 제공
    _HAS_MAIN = True
except Exception:
    rd_main = None
    _HAS_MAIN = False

# ---------------------------
# 공통 유틸 / 메모리
# ---------------------------
def _one_line(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())

def _media_paths():
    media_root = getattr(settings, "MEDIA_ROOT",
                         os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media"))
    media_url  = getattr(settings, "MEDIA_URL", "/media/")
    final_dir  = os.path.join(media_root, "final")
    os.makedirs(final_dir, exist_ok=True)
    mem_dir = os.path.join(final_dir, "_mem")
    os.makedirs(mem_dir, exist_ok=True)
    main_dir = os.path.join(final_dir, "_main")  # main.py 산출물 보관
    os.makedirs(main_dir, exist_ok=True)
    tmp_dir = os.path.join(final_dir, "_tmp")    # 입력 중간물 보관
    os.makedirs(tmp_dir, exist_ok=True)
    return media_root, media_url, final_dir, mem_dir, main_dir, tmp_dir

def _mem_paths(project: Project, user):
    _, _, _, mem_dir, _, _ = _media_paths()
    uid = getattr(user, "id", None) or getattr(user, "pk", None) or "u"
    pid = project.project_id
    last_draft_path = Path(mem_dir) / f"proj{pid}_user{uid}_last_draft.json"
    final_ids_path  = Path(mem_dir) / f"proj{pid}_user{uid}_finalized_ids.json"
    return last_draft_path, final_ids_path

def _remember_last_draft(project: Project, user, draft_id: int):
    last_draft_path, _ = _mem_paths(project, user)
    last_draft_path.write_text(json.dumps({
        "project_id": project.project_id,
        "user_id": getattr(user, "id", None),
        "draft_id": int(draft_id),
        "ts": timezone.now().isoformat()
    }, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"🧠 remembered last draft -> {last_draft_path.as_posix()}")

def _load_last_draft(project: Project, user) -> int | None:
    last_draft_path, _ = _mem_paths(project, user)
    if last_draft_path.exists():
        try:
            obj = json.loads(last_draft_path.read_text(encoding="utf-8"))
            return int(obj.get("draft_id"))
        except Exception:
            return None
    return None

def _remember_finalized_ids(project: Project, user, req_ids: list[int]):
    _, final_ids_path = _mem_paths(project, user)
    final_ids_path.write_text(json.dumps({
        "project_id": project.project_id,
        "user_id": getattr(user, "id", None),
        "requirement_ids": list(map(int, req_ids)),
        "ts": timezone.now().isoformat()
    }, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"🧠 remembered finalized ids -> {final_ids_path.as_posix()}")

def _load_finalized_ids(project: Project, user) -> list[int] | None:
    _, final_ids_path = _mem_paths(project, user)
    if final_ids_path.exists():
        try:
            obj = json.loads(final_ids_path.read_text(encoding="utf-8"))
            ids = obj.get("requirement_ids") or []
            return list(map(int, ids)) if ids else None
        except Exception:
            return None
    return None

def _write_text(path: Path, text: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text or "", encoding="utf-8")

# ---------------------------
# 템플릿
# ---------------------------
def _auto_make_minimal_template(dst: Path) -> Path:
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc = DocxDocument()
    doc.add_heading("연구개발계획서 (자동 생성)", level=0)
    doc.add_heading("Ⅰ. 프로젝트 개요", level=1)
    doc.add_paragraph("프로젝트명: {{프로젝트명}}")
    doc.add_paragraph("생성일시: {{생성일시}}")
    doc.add_heading("Ⅱ. 확정 요구사항 요약", level=1)
    doc.add_paragraph("{{요약_확정요구사항}}")
    doc.add_heading("Ⅲ. 확정 요구사항 본문", level=1)
    doc.add_paragraph("{{원문_확정요구사항}}")
    doc.add_heading("Ⅳ. 선택 Draft 요약", level=1)
    doc.add_paragraph("{{요약_Draft}}")
    doc.add_heading("Ⅴ. 선택 Draft 본문", level=1)
    doc.add_paragraph("{{원문_Draft}}")
    doc.add_heading("Ⅵ. Gemini3 보고서 요약", level=1)
    doc.add_paragraph("{{요약_Gemini3}}")
    doc.add_heading("Ⅶ. Gemini3 보고서 본문", level=1)
    doc.add_paragraph("{{원문_Gemini3}}")
    doc.add_heading("Ⅷ. 섹션 본문(자동 생성)", level=1)
    doc.add_paragraph("{{원문_Main}}")
    doc.save(dst.as_posix())
    return dst

def _default_template_path() -> Path:
    media_root = getattr(settings, "MEDIA_ROOT",
                         os.path.join(getattr(settings, "BASE_DIR", os.getcwd()), "media"))
    base_dir = getattr(settings, "BASE_DIR", os.getcwd())
    for p in [
        Path(media_root) / "templates" / "[연구개발계획서] Potatoes.docx",
        Path(media_root) / "templates" / "potatoes.docx",
        Path(base_dir)  / "templates" / "[연구개발계획서] Potatoes.docx",
        Path(base_dir)  / "templates" / "potatoes.docx",
    ]:
        if p.exists() and p.is_file():
            return p
    return _auto_make_minimal_template(Path(media_root) / "templates" / "auto_default_template.docx")

# ---------------------------
# finalize: 초안 → Requirement 생성/확정
# ---------------------------
def _extract_features_from_draft_content(draft: RequirementDraft):
    raw = (draft.content or "").strip()
    try:
        obj = json.loads(raw)
    except Exception:
        return None
    if isinstance(obj, list):
        return obj
    if isinstance(obj, dict):
        for key in ["정제기획서", "기능목록", "features", "items"]:
            if key in obj and isinstance(obj[key], list):
                return obj[key]
    return None

def _create_requirements_for_features(project: Project, draft: RequirementDraft, features_list):
    created_ids = []
    for item in features_list:
        if isinstance(item, dict):
            fname = item.get("feature_name") or item.get("name") or item.get("기능명") or item.get("title")
            summ  = item.get("summary")      or item.get("desc") or item.get("description") or item.get("설명")
        else:
            fname, summ = None, None

        feature_name = _one_line(fname) if fname else _one_line(str(item)[:120])
        summary      = _one_line(summ) if summ else _one_line(str(item))

        req = Requirement(project=project)
        if hasattr(req, "feature_name"): setattr(req, "feature_name", feature_name)
        if hasattr(req, "summary"):      setattr(req, "summary", summary)
        for flag in ["confirmed_by_user", "is_confirmed", "finalized"]:
            if hasattr(req, flag): setattr(req, flag, True)

        linked = False
        for fk in ["selected_from_draft", "draft", "from_draft", "source_draft"]:
            if hasattr(req, fk):
                setattr(req, fk, draft); linked = True; break
        if not linked:
            did = getattr(draft, "RequirementDraft_id", None) or getattr(draft, "pk", None) or getattr(draft, "id", None)
            for fid in ["draft_id", "from_draft_id", "source_draft_id", "RequirementDraft_id"]:
                if hasattr(req, fid):
                    setattr(req, fid, did); linked = True; break

        req.save()
        pk = getattr(req, "pk", None) or getattr(req, "id", None)
        created_ids.append(int(pk))
    return created_ids

class FinalizeRequirementView(APIView):
    """
    POST /api/project/<project_id>/finalize/
    Body(JSON): { "draft_id": 45 }
    """
    permission_classes = [IsAuthenticated]

    def post(self, request, project_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)
        try:
            draft_id = int(request.data.get("draft_id"))
        except Exception:
            return Response({"error": "draft_id는 필수입니다.(정수)"}, status=400)

        draft = get_object_or_404(RequirementDraft, pk=draft_id, project=project)
        features = _extract_features_from_draft_content(draft)
        if not features:
            return Response({"error": "초안에서 기능 목록을 파싱하지 못했습니다."}, status=400)

        try:
            created_ids = _create_requirements_for_features(project, draft, features)
        except Exception as e:
            return Response({"error": f"Requirement 생성 실패: {e}"}, status=500)

        _remember_last_draft(project, request.user, draft_id)
        _remember_finalized_ids(project, request.user, created_ids)

        print(f"✅ finalize: project={project.project_id}, draft_id={draft_id}, count={len(created_ids)}")
        return Response({"message": "최종 기능 명세 저장 완료", "count": len(created_ids), "requirement_ids": created_ids}, status=201)

# ---------------------------
# generate: 최종 개발문서 생성
# ---------------------------
def _requirements_for_draft(project: Project, draft: RequirementDraft):
    """선택 draft로 생성/확정된 Requirement만 찾는다. (order_by -> pk 로 통일)"""
    draft_pk = getattr(draft, "RequirementDraft_id", None) or getattr(draft, "pk", None) or getattr(draft, "id", None)

    qs = None
    # 1) FK 필드들 우선 시도
    for fname in ["selected_from_draft", "draft", "from_draft", "source_draft"]:
        try:
            k = {fname: draft}
            qs_try = Requirement.objects.filter(project=project, **k).order_by("pk")  # ★ fixed
            if qs_try.exists():
                qs = list(qs_try); break
        except FieldError:
            continue

    # 2) 정수 ID 필드들 시도
    if not qs:
        for fname in ["draft_id", "from_draft_id", "source_draft_id", "RequirementDraft_id"]:
            try:
                k = {fname: draft_pk}
                qs_try = Requirement.objects.filter(project=project, **k).order_by("pk")  # ★ fixed
                if qs_try.exists():
                    qs = list(qs_try); break
            except FieldError:
                continue

    if not qs:
        return None, "선택한 초안으로 생성/확정된 Requirement를 찾지 못했습니다. (연결 컬럼 확인 필요)"

    bullets = [f"- {_one_line(r.feature_name)}: {_one_line(getattr(r, 'summary', '') or '')}" for r in qs]
    req_summary = "\n".join(bullets)[:4000]
    req_full    = "\n".join([f"{r.feature_name}\n{getattr(r, 'summary', '') or ''}" for r in qs])[:12000]
    return {"요약_확정요구사항": req_summary or "(요약 없음)", "원문_확정요구사항": req_full or "(본문 없음)"}, None

def _collect_3_sources_strict(project: Project, draft: RequirementDraft, user):
    """
    1) 선택 Draft
    2) 그 Draft로 생성된 Requirement 집합 (링크 컬럼 → 실패 시 finalize 메모리 PK 폴백)
    3) Gemini3 리포트 MD (필수)
    """
    # 1) Draft 본문/요약
    draft_raw = (draft.content or "").strip()
    try:
        draft_snippet = json.dumps(json.loads(draft_raw), ensure_ascii=False, indent=2)[:4000]
    except Exception:
        draft_snippet = draft_raw[:4000]
    draft_summary = _one_line(getattr(draft, "summary", "") or "") or "(요약 없음)"

    # 2) Requirement 수집 (링크 컬럼)
    req_map, err = _requirements_for_draft(project, draft)
    if err:
        ids = _load_finalized_ids(project, user)
        if ids:
            qs_try = Requirement.objects.filter(project=project, pk__in=ids).order_by("pk")  # ★ fixed
            if qs_try.exists():
                reqs = list(qs_try)
                bullets = [f"- {_one_line(r.feature_name)}: {_one_line(getattr(r, 'summary', '') or '')}" for r in reqs]
                req_summary = "\n".join(bullets)[:4000]
                req_full    = "\n".join([f"{r.feature_name}\n{getattr(r, 'summary', '') or ''}" for r in reqs])[:12000]
                req_map = {"요약_확정요구사항": req_summary or "(요약 없음)", "원문_확정요구사항": req_full or "(본문 없음)"}
                err = None
    if err:
        return None, {"error": err}

    # 3) Gemini3 리포트
    media_root, _, _, _, _, _ = _media_paths()
    report_fs = os.path.join(media_root, "reports", f"project{project.project_id}_analysis_report.md")
    if not os.path.exists(report_fs):
        return None, {"error": f"Gemini3 보고서를 찾지 못했습니다: {report_fs}"}
    try:
        gem3_md = Path(report_fs).read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return None, {"error": "Gemini3 보고서 읽기 실패"}

    mapping = {
        "프로젝트명": project.title or f"project {project.project_id}",
        "생성일시": timezone.now().strftime("%Y-%m-%d %H:%M"),
        **req_map,
        "요약_Draft": draft_summary,
        "원문_Draft": draft_snippet or "",
        "요약_Gemini3": (gem3_md.splitlines()[0] if gem3_md else "유사도 보고서 없음"),
        "원문_Gemini3": gem3_md or "",
    }
    return mapping, None

# ---------------------------
# main.py 파이프라인: 입력 준비 & 실행
# ---------------------------
def _build_main_inputs(project: Project, mapping: dict, tmp_dir: str) -> list[Path]:
    """
    main.run_pipeline 에 넣을 근거 파일들을 Markdown 으로 구성
    - plan.md : Project.description
    - spec.md : 확정된 Requirement 요약/본문
    - report.md : Gemini3 보고서 전문
    - draft.md : 선택 Draft 전문(요약은 상단에 주석)
    """
    pid = project.project_id
    p_plan   = Path(tmp_dir) / f"project{pid}_plan.md"
    p_spec   = Path(tmp_dir) / f"project{pid}_spec.md"
    p_report = Path(tmp_dir) / f"project{pid}_gemini3_report.md"
    p_draft  = Path(tmp_dir) / f"project{pid}_draft.md"

    # 1) plan
    plan_md = (project.description or "").strip()
    if not plan_md:
        plan_md = "# 프로젝트 개요\n(프로젝트 설명이 비어 있음)"
    _write_text(p_plan, plan_md)

    # 2) spec (요약 + 본문)
    spec_md = "## 확정 요구사항 요약\n\n" + (mapping.get("요약_확정요구사항") or "")
    spec_md += "\n\n## 확정 요구사항 본문\n\n" + (mapping.get("원문_확정요구사항") or "")
    _write_text(p_spec, spec_md)

    # 3) report
    _write_text(p_report, mapping.get("원문_Gemini3") or "# Gemini3 보고서 없음")

    # 4) draft
    draft_head = mapping.get("요약_Draft") or ""
    draft_body = mapping.get("원문_Draft") or ""
    _write_text(p_draft, f"<!-- 요약: {draft_head} -->\n\n{draft_body}")

    return [p_plan, p_spec, p_report, p_draft]

def _run_main_pipeline(project: Project, out_docx_path: Path, inputs: list[Path]) -> tuple[Path | None, str | None]:
    if not _HAS_MAIN:
        return None, "main.py 모듈을 찾지 못했습니다. (의존 모듈 src.* 확인 필요)"
    # API 키: GOOGLE_API_KEY 또는 GEMINI_API_KEY
    api_key = os.environ.get("GOOGLE_API_KEY") or os.environ.get("GEMINI_API_KEY") or os.environ.get("GEMINI_API_KEY_3")
    if not api_key:
        return None, "환경변수 GOOGLE_API_KEY 또는 GEMINI_API_KEY 가 필요합니다."

    # ✅ plan.json 경로를 명시적으로 지정(빈 문자열 금지)
    plan_json_path = out_docx_path.parent / f"proj{project.project_id}_plan.json"
    plan_json_path.parent.mkdir(parents=True, exist_ok=True)

    args = types.SimpleNamespace(
        inputs=[p.as_posix() for p in inputs],
        plan=plan_json_path.as_posix(),                 # ← 수정: 실제 파일 경로 전달
        out=out_docx_path.as_posix(),
        model=os.environ.get("MAIN_MODEL_NAME", "gemini-2.5-flash"),
        api_key=api_key,
        max_chunk_tokens=int(os.environ.get("MAIN_MAX_CHUNK_TOKENS", "1200")),
        max_retry=int(os.environ.get("MAIN_MAX_RETRY", "3")),
        temp=float(os.environ.get("MAIN_TEMPERATURE", "0.2")),
    )
    try:
        rd_main.run_pipeline(args)
    except Exception as e:
        return None, f"main.run_pipeline 실패: {e}"
    if not out_docx_path.exists():
        return None, "main.run_pipeline 산출물이 존재하지 않습니다."
    return out_docx_path, None

# ---------------------------
# View
# ---------------------------
class FinalDevDocGenerateView(APIView):
    """
    POST /api/project/<project_id>/final-devdoc/generate/
      - Body JSON: {"draft_id": 45} (옵션)
      - 또는 쿼리스트링: ?draft_id=45 (옵션)
      - 아무 것도 없으면 '최근 사용 draft' 기본값 사용
    결과: media/final/project{pid}_{ts}_final.docx 저장 + URL 반환
          (추가) main.py 섹션 문서 별도 저장 URL도 함께 반환
    """
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request, project_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)

        # draft_id: body > query > last-used
        raw = request.data.get("draft_id", "") or request.query_params.get("draft_id")
        if raw in (None, ""):
            raw = _load_last_draft(project, request.user)
        try:
            draft_id = int(raw)
        except Exception:
            return Response({"error": "draft_id가 없습니다. 최근 선택값도 기억되어 있지 않습니다."}, status=400)

        draft = get_object_or_404(RequirementDraft, pk=draft_id, project=project)
        _remember_last_draft(project, request.user, draft_id)  # 최신값 다시 기억

        # 템플릿
        tpl_path = _default_template_path()
        # 치환 데이터(3개 소스)
        mapping, err = _collect_3_sources_strict(project, draft, request.user)
        if err:
            return Response(err, status=400)

        # 저장 경로들
        media_root, media_url, final_dir, _mem, main_dir, tmp_dir = _media_paths()
        ts = timezone.now().strftime("%Y%m%d_%H%M%S")
        final_dst = Path(final_dir) / f"project{project.project_id}_{ts}_final.docx"
        main_dst  = Path(main_dir)  / f"project{project.project_id}_{ts}_main.docx"

        # ========== 1) main.py 파이프라인 실행 (섹션 산출) ==========
        main_text = ""
        main_file_url = None
        try:
            inputs = _build_main_inputs(project, mapping, tmp_dir)
            outp, err2 = _run_main_pipeline(project, main_dst, inputs)
            if err2:
                print(f"[경고] main.py 파이프라인 생략: {err2}")
            else:
                # main.docx → 텍스트 추출 (템플릿 치환용)
                _, main_text = load_docx_and_plaintext(outp)
                # URL
                main_file_url = f"{media_url}final/_main/{main_dst.name}"
                # 템플릿에 키가 있다면 매핑에 추가
                mapping.setdefault("원문_Main", (main_text or "")[:50000])
        except Exception as e:
            print(f"[경고] main.py 통합 처리 중 예외: {e}")

        # ========== 2) 템플릿 DOCX 생성 ==========
        try:
            tpl_doc, tpl_text = load_docx_and_plaintext(tpl_path)
            has_main_placeholder = "{{원문_Main}}" in tpl_text

            doc = DocxDocument(tpl_path.as_posix())
            docx_replace_placeholders(doc, mapping, unsure_to_red=False)

            # 템플릿에 {{원문_Main}} 자리가 없으면 '부록'으로 덧붙임
            if (not has_main_placeholder) and main_text:
                doc.add_page_break()
                doc.add_heading("부록. 자동 생성 섹션 본문(main.py)", level=1)
                for line in (main_text or "").splitlines():
                    doc.add_paragraph(line)

            doc.save(final_dst.as_posix())
        except Exception as e:
            return Response({"error": f"DOCX 생성 실패: {e}"}, status=500)

        file_url = f"{media_url}final/{final_dst.name}"
        print(f"📄 최종 개발문서 저장: {final_dst.as_posix()}")
        if main_file_url:
            print(f"📄 섹션 본문(main) 저장: {main_dst.as_posix()}")

        return Response({
            "ok": True,
            "file_url": file_url,
            "main_file_url": main_file_url,  # 없을 수도 있음
            "draft_id_used": draft_id
        }, status=201)





# 파일 상단에 추가
import glob

class FinalDevDocFilesView(APIView):
    """
    GET /api/project/<project_id>/final-devdoc/files/
    설명: media/final/에 저장된 해당 프로젝트의 최종 문서 목록을 최신순으로 반환
    옵션: ?include_main=1 로 보내면 main.py 산출물(_main)도 함께 반환
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, project_id):
        project = get_object_or_404(Project, pk=project_id, user=request.user)

        # ✅ _media_paths가 6개를 반환하므로 이렇게 받아야 함
        _, media_url, final_dir, *_ = _media_paths()

        # 최종(final) 문서
        pattern = os.path.join(final_dir, f"project{project.project_id}_*_final.docx")
        files = sorted(glob.glob(pattern), key=os.path.getmtime, reverse=True)

        items = []
        for fp in files:
            name = os.path.basename(fp)
            # 타임스탬프 추출: project{pid}_YYYYMMDD_HHMMSS_final.docx
            m = re.search(rf"project{project.project_id}_(\d{{8}}_\d{{6}})_final\.docx$", name)
            ts = m.group(1) if m else ""
            items.append({
                "file_url": f"{media_url}final/{name}",
                "filename": name,
                "created_at": ts
            })

        resp = {"items": items, "latest": (items[0] if items else None)}

        # (옵션) main.py 결과도 함께
        include_main = str(request.query_params.get("include_main", "")).lower() in ("1", "true", "yes")
        if include_main:
            main_dir = os.path.join(final_dir, "_main")
            main_pattern = os.path.join(main_dir, f"project{project.project_id}_*_main.docx")
            main_files = sorted(glob.glob(main_pattern), key=os.path.getmtime, reverse=True)

            main_items = []
            for fp in main_files:
                name = os.path.basename(fp)
                m = re.search(rf"project{project.project_id}_(\d{{8}}_\d{{6}})_main\.docx$", name)
                ts = m.group(1) if m else ""
                main_items.append({
                    "file_url": f"{media_url}final/_main/{name}",
                    "filename": name,
                    "created_at": ts
                })

            resp["main_items"] = main_items
            resp["main_latest"] = main_items[0] if main_items else None

        print(f"📂 최종 문서 목록 조회(project {project.project_id}): {len(items)}건")
        return Response(resp, status=200)



#챗봇
# --- Chatbot proxy: views -> chat.py를 통해 Gemini 호출 -----------------------
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from django.conf import settings
import os

# 1) chat.py를 불러와서 제공 함수/상수 우선 사용
try:
    from . import chat as chatmod
except Exception as e:
    chatmod = None

# 2) chat.py가 없거나 일부 함수가 없을 때를 대비한 안전 폴백들
import google.generativeai as genai
from google.generativeai.types import GenerationConfig

def _get_api_key():
    # chat.py가 제공하면 우선 사용
    if chatmod and hasattr(chatmod, "get_api_key"):
        key = chatmod.get_api_key()
        if key:
            return key
    # 우선순위: GEMINI_API_KEY_4 > GEMINI_API_KEY > settings.GEMINI_API_KEY
    return (
        os.getenv("GEMINI_API_KEY_4")
        or os.getenv("GEMINI_API_KEY")
        or getattr(settings, "GEMINI_API_KEY", None)
    )

def _configure_gemini():
    api_key = _get_api_key()
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY가 설정되어 있지 않습니다.")
    genai.configure(api_key=api_key)

def _get_system_instruction():
    # chat.py가 시스템 지침 제공 시 사용
    if chatmod and hasattr(chatmod, "SYSTEM_INSTRUCTION_KO"):
        return getattr(chatmod, "SYSTEM_INSTRUCTION_KO")
    if chatmod and hasattr(chatmod, "SYSTEM_PROMPT"):
        return getattr(chatmod, "SYSTEM_PROMPT")
    # 기본 폴백
    return (
        "역할: 제품/기능 기획 코파일럿. "
        "원칙: 1) 의도확인 2) 요약 3) 선택지 제시 4) 근거 5) 다음 액션 제안. "
        "출력: 한국어."
    )

def _build_model():
    # chat.py가 모델 팩토리/래퍼 제공 시 사용
    if chatmod and hasattr(chatmod, "get_model"):
        return chatmod.get_model()
    if chatmod and hasattr(chatmod, "init_model"):
        return chatmod.init_model()
    # 기본 Gemini 모델 생성
    _configure_gemini()
    return genai.GenerativeModel(
        model_name=getattr(chatmod, "MODEL_NAME", "gemini-1.5-flash"),
        system_instruction=_get_system_instruction(),
    )

def _to_gemini_history(raw_history):
    """
    raw_history: [{"role":"user"|"model","content":"..."}...]
    chat.py가 히스토리 변환기를 제공하면 그걸 사용.
    """
    if chatmod and hasattr(chatmod, "to_gemini_history"):
        try:
            return chatmod.to_gemini_history(raw_history or [])
        except Exception:
            pass
    hist = []
    for h in (raw_history or []):
        if not isinstance(h, dict):
            continue
        role = h.get("role", "user")
        content = h.get("content", "")
        if content is None:
            continue
        hist.append({"role": role, "parts": [str(content)]})
    return hist

def _postprocess_reply(text: str) -> str:
    # chat.py가 후처리를 제공하면 우선 사용
    if chatmod and hasattr(chatmod, "postprocess_reply"):
        try:
            return chatmod.postprocess_reply(text)
        except Exception:
            pass
    return text


class ChatbotView(APIView):
    """
    POST /api/chat/
    body:
      {
        "message": "질문/명령",            # 필수
        "history": [                      # 선택
          {"role":"user","content":"..."},
          {"role":"model","content":"..."}
        ],
        "mime": "text/plain",             # 선택 (기본 text/plain)
        "model": "gemini-1.5-flash"       # 선택 (chat.py에서 무시/재정의 가능)
      }
    """
    permission_classes = [IsAuthenticated]

    def post(self, request):
        msg = (request.data.get("message") or "").strip()
        if not msg:
            return Response({"error": "message는 필수입니다."}, status=400)

        raw_history = request.data.get("history") or []
        mime = (request.data.get("mime") or "text/plain").strip()
        preferred_model = (request.data.get("model") or "").strip()

        try:
            model = None

            # chat.py가 통합 send 함수를 제공한다면 최우선 사용
            if chatmod and hasattr(chatmod, "send_chat"):
                try:
                    reply, usage = chatmod.send_chat(
                        message=msg,
                        history=raw_history,
                        mime=mime,
                        model_name=preferred_model or None
                    )
                    return Response({
                        "reply": reply,
                        "usage": usage,
                        "source": "chat.py:send_chat"
                    }, status=200)
                except Exception as e:
                    # 실패하면 표준 경로로 폴백
                    pass

            # 표준 경로: 모델 생성 → 히스토리 세팅 → 메시지 전송
            model = _build_model()

            # 모델 이름 오버라이드가 필요하면 교체(옵션)
            if preferred_model:
                _configure_gemini()
                model = genai.GenerativeModel(
                    model_name=preferred_model,
                    system_instruction=_get_system_instruction(),
                )

            history = _to_gemini_history(raw_history)
            chat = model.start_chat(history=history)

            resp = chat.send_message(
                msg,
                generation_config=GenerationConfig(
                    temperature=getattr(chatmod, "TEMPERATURE", 0.3),
                    candidate_count=1,
                    max_output_tokens=getattr(chatmod, "MAX_OUTPUT_TOKENS", 1024),
                    response_mime_type=mime
                )
            )

            text = _postprocess_reply(getattr(resp, "text", "") or "")
            usage = getattr(resp, "usage_metadata", None)
            usage_payload = {
                "prompt_tokens": getattr(usage, "prompt_token_count", None),
                "candidates_tokens": getattr(usage, "candidates_token_count", None),
                "total_tokens": getattr(usage, "total_token_count", None),
            } if usage else None

            return Response({
                "reply": text,
                "usage": usage_payload,
                "source": "views+fallback"
            }, status=200)

        except Exception as e:
            msg = str(e).lower()
            if "429" in msg or "quota" in msg or "rate" in msg:
                return Response({
                    "error": "요청이 많아 일시적으로 제한되었습니다(HTTP 429).",
                    "hint": "Gemini 쿼터/레이트 리밋. API 키/쿼터 확인 후 재시도."
                }, status=429)
            if "401" in msg or "unauthorized" in msg or "invalid" in msg:
                return Response({
                    "error": "인증 실패(HTTP 401).",
                    "hint": "GEMINI_API_KEY 환경변수 또는 chat.py 키 설정 확인"
                }, status=401)
            return Response({"error": f"챗봇 처리 중 오류: {e}"}, status=500)

# === Latest Gantt Tasks View (화면 렌더링용) ===
from django.shortcuts import get_object_or_404
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions, status

# 모델 import (이미 위에서 import 했다면 중복 제거)
from .models import Project, GanttChart, GanttTask


class LatestGanttTasksView(APIView):
    """
    GET /api/project/<int:project_id>/gantt/latest/tasks/
    - 해당 프로젝트의 '가장 최근' 간트차트를 찾아 태스크 목록을 JSON으로 반환
    - 프론트는 이 JSON으로 화면에 간트를 렌더링하고,
      별도의 다운로드 URL(생성 응답 or 기존 다운로드 뷰)을 버튼에 연결하면 됨.
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, project_id: int):
        # 1) 프로젝트 권한 체크 (해당 유저의 프로젝트만)
        project = get_object_or_404(Project, project_id=project_id, user=request.user)

        # 2) 최신 간트차트 1개 선택
        latest = (
            GanttChart.objects
            .filter(project=project)
            .order_by("-created_at", "-GanttChart")  # created_at 우선, 동시 생성시 PK 내림차순
            .first()
        )
        if not latest:
            return Response({"error": "간트차트가 없습니다."}, status=status.HTTP_404_NOT_FOUND)

        # 3) 태스크 조회 (파트/시작주차/순서 기준 정렬)
        tasks_qs = (
            GanttTask.objects
            .filter(gantt_chart=latest)
            .select_related("requirement")
            .order_by("part", "start_week", "order", "GanttTask")  # ← 'id' 대신 'GanttTask'
        )

        tasks = []
        for t in tasks_qs:
            tasks.append({
                "id": getattr(t, "GanttTask", getattr(t, "pk", None)),  # ← PK 필드명 반영
                "part": t.part,
                "feature_name": t.feature_name,
                "start_week": t.start_week,
                "duration_weeks": t.duration_weeks,
                "order": t.order,
                "requirement_id": (
                    getattr(t.requirement, "Requirement", getattr(t.requirement, "pk", None))
                    if t.requirement else None
                ),
            })

        # 4) 화면 렌더용 메타 + 태스크 반환
        return Response({
            "gantt_id": getattr(latest, "GanttChart", getattr(latest, "pk", None)),
            "project_id": getattr(project, "project_id", None),
            "start_date": latest.start_date,     # e.g. "2025-08-19"
            "total_weeks": latest.total_weeks,   # e.g. 12
            "parts": latest.parts,               # e.g. ["백엔드","프론트엔드","AI","서류"]
            "tasks": tasks,                      # 위에서 구성한 태스크 배열
        }, status=status.HTTP_200_OK)

# views.py (필요 부분만)
import os
from django.conf import settings
from django.http import Http404
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions
from .idea_refiner import IdeaRefiner
from .idea_expander import IdeaExpander



class IdeaPreviewView(APIView):
    """
    GET /api/idea/preview/?idea=...&title=...&also_return_markdown=true
    GET /api/idea/preview/?filename=idea_plan_1.md
    - 저장 없이 정제/확장(MD 포함) OR 기존 md 파일을 읽어 반환
    """
    permission_classes = [permissions.IsAuthenticated]

    def _fallback_md(self, refined, suggestions):
        lines = [
            "# 기획서 미리보기",
            f"## 목표\n{refined.get('goal','')}",
            f"## 대상 사용자\n{refined.get('target_user','')}",
        ]
        core = refined.get("core_features") or []
        if core:
            lines.append("## 핵심 기능")
            for i, f in enumerate(core, 1):
                lines.append(f"{i}. {f}")
        if suggestions:
            lines.append("## 제안/추천")
            for s in suggestions:
                lines.append(f"- {s}")
        return "\n\n".join(lines)

    def get(self, request, *args, **kwargs):
        filename = request.query_params.get("filename")
        idea = request.query_params.get("idea")
        title = request.query_params.get("title") or ""
        want_md = str(request.query_params.get("also_return_markdown", "true")).lower() in ("1","true","yes")

        # 1) 파일명으로 미리보기 (이미 export된 md 읽기)
        if filename:
            base = getattr(settings, "MEDIA_ROOT", os.path.join(settings.BASE_DIR, "media"))
            md_path = os.path.join(base, "ideas", filename)
            if not (os.path.isfile(md_path) and md_path.endswith(".md")):
                raise Http404("markdown file not found")
            with open(md_path, "r", encoding="utf-8") as f:
                md_text = f.read()
            return Response({"md": md_text, "from_file": True}, status=200)

        # 2) 아이디어로 즉시 정제/확장 (저장 없음)
        if not idea:
            return Response({"error": "idea or filename is required"}, status=400)

        refined = IdeaRefiner().refine(idea)
        suggestions = IdeaExpander().expand(refined)
        similar_map = {}  # 필요 시 채워 넣기

        md_text = None
        if want_md:
            try:
                md_text = generate_markdown(refined, suggestions, similar_map)
            except Exception:
                md_text = self._fallback_md(refined, suggestions)

        payload = {
            "refined": refined,
            "suggestions": suggestions,
            "similar_map": similar_map,
        }
        if want_md:
            payload["md"] = md_text
        return Response(payload, status=200)



# views.py
import os, mimetypes
from django.conf import settings
from django.http import FileResponse, Http404
from rest_framework.views import APIView
from rest_framework import permissions

class IdeaFileDownloadView(APIView):
    """
    생성된 아이디어 기획서(MD/DOCX) 파일 다운로드
    GET /api/idea/download/<filename>/
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, filename, *args, **kwargs):
        base = getattr(settings, "MEDIA_ROOT", os.path.join(settings.BASE_DIR, "media"))
        ideas_dir = os.path.abspath(os.path.join(base, "ideas"))
        filepath = os.path.abspath(os.path.join(ideas_dir, filename))

        # 경로 탈출 방지 + 존재 확인
        if not filepath.startswith(ideas_dir) or not os.path.exists(filepath):
            raise Http404("File not found")

        content_type, _ = mimetypes.guess_type(filepath)
        resp = FileResponse(open(filepath, "rb"), as_attachment=True, filename=filename)
        if content_type:
            resp["Content-Type"] = content_type
        return resp

