# -*- coding: utf-8 -*-
import glob
import time
import json
import os
import requests
from dotenv import load_dotenv

# DRF 뷰에서는 아래 세 함수만 사용합니다.
# 콘솔용(정제/확장/대화) 클래스들은 main() 내부에서만 임포트합니다.

from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

import google.generativeai as genai

# ===================== 환경변수 / 상수 =====================
ENV_KEY_NAME = "GEMINI_API_KEY_4"   # 이 키로 Gemini API Key를 읽습니다.
_GEMINI_MODEL = None                # 지연 초기화용 캐시
_GITHUB_WARNED = False              # 토큰 경고 1회만 표시

# GitHub 검색 상한
MAX_KEYWORDS_PER_FEATURE = 2
MAX_REPOS_PER_FEATURE = 3
GITHUB_PER_PAGE = MAX_REPOS_PER_FEATURE * 2  # 중복 제거 전 여유 수집


def _get_gemini_model():
    """
    Gemini 모델 지연 초기화.
    - 모듈 import 시점에 환경변수 없어서 서버가 죽는 문제를 방지
    - 실제 호출 시점에만 환경변수 로드/검증
    """
    global _GEMINI_MODEL
    if _GEMINI_MODEL is not None:
        return _GEMINI_MODEL

    load_dotenv()
    api_key = os.getenv(ENV_KEY_NAME)
    if not api_key:
        raise RuntimeError(
            f"{ENV_KEY_NAME} 환경 변수가 없습니다. "
            f".env 또는 서버 환경에 {ENV_KEY_NAME}=<YOUR_KEY> 를 설정하세요."
        )
    genai.configure(api_key=api_key)
    _GEMINI_MODEL = genai.GenerativeModel('gemini-1.5-flash')
    return _GEMINI_MODEL


def _get_github_token():
    """
    GitHub 토큰 지연 로드(없어도 동작하지만 rate limit이 낮습니다).
    """
    global _GITHUB_WARNED
    load_dotenv()
    tok = os.getenv("GITHUB_TOKEN")
    if not tok and not _GITHUB_WARNED:
        print("⚠ 경고: GITHUB_TOKEN이 없어 GitHub 검색 쿼터가 제한됩니다.")
        _GITHUB_WARNED = True
    return tok


# ===================== 공통 유틸 =====================
def gemini_call_with_retry(prompt: str, max_retries: int = 3) -> str:
    """Gemini 호출 with 재시도 (지연 초기화 포함)"""
    model = _get_gemini_model()
    for attempt in range(max_retries):
        try:
            resp = model.generate_content(prompt)
            return (resp.text or "").strip()
        except Exception as e:
            print(f"❌ Gemini 호출 실패: {e}")
            if attempt < max_retries - 1:
                print(f"⏳ {attempt + 1}회 실패 → 재시도 중...")
                time.sleep(2)
            else:
                raise


def github_search_repos(query: str, per_page: int = GITHUB_PER_PAGE):
    """GitHub 저장소 검색 → dict 리스트 반환(name, full_name, url, stars, desc)."""
    url = "https://api.github.com/search/repositories"
    headers = {"Accept": "application/vnd.github+json"}
    tok = _get_github_token()
    if tok:
        headers["Authorization"] = f"Bearer {tok}"
    params = {"q": query, "sort": "stars", "order": "desc", "per_page": per_page}
    r = requests.get(url, headers=headers, params=params, timeout=10)
    r.raise_for_status()
    items = r.json().get("items", [])
    out = []
    for it in items:
        out.append({
            "name": it.get("name", ""),
            "full_name": it.get("full_name") or it.get("name", ""),
            "url": it.get("html_url", ""),
            "stars": it.get("stargazers_count", 0),
            "desc": (it.get("description") or "").strip()
        })
    return out


def _flatten_to_strings(maybe_nested):
    """리스트/딕셔너리/문자열 등이 섞인 구조를 납작한 문자열 리스트로 변환."""
    out = []
    def walk(x):
        if x is None:
            return
        if isinstance(x, str):
            s = x.strip()
            if s:
                out.append(s)
        elif isinstance(x, (list, tuple, set)):
            for it in x:
                walk(it)
        elif isinstance(x, dict):
            for v in x.values():
                walk(v)
        else:
            out.append(str(x))
    walk(maybe_nested)
    return out


def _json_from_gemini_text(resp_text: str):
    """```json ... ``` 래퍼 제거 + JSON 파싱"""
    t = (resp_text or "").strip()
    if t.startswith("```json"):
        t = t[7:]
    if t.endswith("```"):
        t = t[:-3]
    return json.loads(t or "[]")


# ===================== 유사 기능 자동 검색 (Gemini + GitHub) =====================
def build_similar_map(core_features):
    """
    반환:
    {
      "<주요 기능 문장>": [
        {"keyword": "<kw>", "repos": [ {name, full_name, url, stars, desc}, ... ]},
        ...
      ],
      ...
    }
    """
    result = {f: [] for f in core_features}
    chunk_size = 2

    # 1) 기능별 키워드 추출 (Gemini)
    feature_keywords = {}
    for i in range(0, len(core_features), chunk_size):
        chunk = core_features[i:i+chunk_size]
        prompt = f"""
        다음 기능 문장들 각각에 대해 검색에 유용한 핵심 키워드 1~3개를 뽑아
        기능문장을 key로 하는 JSON 객체로만 출력하세요.
        예:
        {{"GPS 기반 방문 인증": ["geofencing","GPS check-in"]}}
        기능:
        {chunk}
        """
        try:
            data = _json_from_gemini_text(gemini_call_with_retry(prompt))
            for feat, kws in (data or {}).items():
                flat = _flatten_to_strings(kws)
                clean = []
                for k in flat:
                    k = k.strip()
                    if len(k) >= 2 and k.lower() not in [x.lower() for x in clean]:
                        clean.append(k)
                feature_keywords[feat] = clean[:MAX_KEYWORDS_PER_FEATURE]
        except Exception as e:
            print(f"⚠ Gemini 키워드 추출 실패: {e}")

    # 2) 키워드로 GitHub 검색 (상위 N개만), 중복 제거
    for feat in core_features:
        entries = []
        seen_fullnames = set()
        for kw in feature_keywords.get(feat, []):
            try:
                repos = github_search_repos(kw, per_page=GITHUB_PER_PAGE)
                curated = []
                for r in repos:
                    if r["full_name"] in seen_fullnames:
                        continue
                    seen_fullnames.add(r["full_name"])
                    # 설명 요약
                    r["desc"] = (r["desc"][:100] + "…") if len(r["desc"]) > 100 else r["desc"]
                    curated.append(r)
                    if len(curated) >= MAX_REPOS_PER_FEATURE:
                        break
                if curated:
                    entries.append({"keyword": kw, "repos": curated})
            except Exception as e:
                print(f"⚠ GitHub 검색 실패({kw}): {e}")

        # 키워드가 비었거나 결과가 없으면 기능 문장으로 보정 검색 1회
        if not entries:
            fallback_kw = " ".join([w for w in feat.split() if len(w) >= 2][:3])
            if fallback_kw:
                try:
                    repos = github_search_repos(fallback_kw, per_page=GITHUB_PER_PAGE)
                    curated = []
                    seen_fullnames = set()
                    for r in repos:
                        if r["full_name"] in seen_fullnames:
                            continue
                        seen_fullnames.add(r["full_name"])
                        r["desc"] = (r["desc"][:100] + "…") if len(r["desc"]) > 100 else r["desc"]
                        curated.append(r)
                        if len(curated) >= MAX_REPOS_PER_FEATURE:
                            break
                    if curated:
                        entries.append({"keyword": fallback_kw, "repos": curated})
                except Exception as e:
                    print(f"⚠ GitHub 보정 검색 실패({fallback_kw}): {e}")

        result[feat] = entries
    return result


# ===================== Markdown 생성 (표/링크/별점) =====================
def generate_markdown(idea_data, suggestions, similar_map):
    md = []
    md.append("# 프로젝트 기획서\n")
    md.append(f"## 🎯 목표\n{idea_data['goal']}\n")
    md.append(f"## 👥 타겟 사용자\n{idea_data['target_user']}\n")

    md.append("## ✨ 주요 기능\n")
    for f in idea_data['core_features']:
        md.append(f"- {f}\n")

    md.append("\n## 🔵 유사 기능 (실제 프로젝트 사례)\n")
    for feat, entries in similar_map.items():
        md.append(f"### <span style='color:blue'>{feat}</span>\n")
        if not entries:
            md.append("- (유사 사례를 찾지 못했습니다)\n\n")
            continue
        # 표 헤더
        md.append("| 프로젝트 | ⭐ Stars | 설명 |\n")
        md.append("|---|---:|---|\n")
        for e in entries:
            kw = e.get("keyword", "")
            if kw:
                md.append(f"> 키워드: `{kw}`\n\n")
            for r in e.get("repos", []):
                name = r["full_name"]
                url = r["url"]
                stars = r["stars"]
                desc = r["desc"].replace("\n", " ")
                md.append(f"| [{name}]({url}) | {stars} | {desc or '-'} |\n")
        md.append("\n")

    md.append("## 🔴 확장 아이디어\n")
    for s in suggestions:
        md.append(f"- <span style='color:red'>{s}</span>\n")

    return "".join(md)


# ===================== Word 생성 (하이퍼링크/별점) =====================
def _add_hyperlink(paragraph, url, text, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); r_pr.append(u)
    if bold:
        b = OxmlElement('w:b'); r_pr.append(b)
    new_run.append(r_pr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def generate_word(idea_data, suggestions, similar_map, file_name):
    doc = Document()
    doc.add_heading("프로젝트 기획서", level=1)

    doc.add_heading("🎯 목표", level=2); doc.add_paragraph(idea_data['goal'])
    doc.add_heading("👥 타겟 사용자", level=2); doc.add_paragraph(idea_data['target_user'])

    doc.add_heading("✨ 주요 기능", level=2)
    for f in idea_data['core_features']:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading("🔵 유사 기능 (실제 프로젝트 사례)", level=2)
    for feat, entries in similar_map.items():
        p = doc.add_paragraph()
        run = p.add_run(feat); run.font.color.rgb = RGBColor(0, 0, 255); run.bold = True

        if not entries:
            doc.add_paragraph("(유사 사례를 찾지 못했습니다)", style='List 2')
            continue

        for e in entries:
            kw = e.get("keyword", "")
            if kw:
                doc.add_paragraph(f"키워드: {kw}", style='List 2')
            for r in e.get("repos", []):
                line = doc.add_paragraph(style='List 3')
                _add_hyperlink(line, r["url"], r["full_name"], bold=True)
                line.add_run(f"  ⭐ {r['stars']}  —  ")
                desc = r["desc"] or "-"
                line.add_run(desc)

    doc.add_heading("🔴 확장 아이디어", level=2)
    for s in suggestions:
        p = doc.add_paragraph(style='List Bullet')
        t = p.add_run(s); t.font.color.rgb = RGBColor(255, 0, 0)

    doc.save(file_name)


# ===================== 콘솔 데모 (선택) =====================
def main():
    """
    콘솔에서만 사용하는 데모 흐름.
    Django/DRF 서버에서는 사용하지 않습니다.
    """
    # 콘솔 모드에서만 필요한 의존성들을 여기서 임포트
    from .idea_refiner import IdeaRefiner
    from .idea_expander import IdeaExpander
    from .conversation_manager import ConversationManager

    print("AI: 안녕하세요! 머릿속에 있는 아이디어를 들려주세요. 제가 구체적인 기획으로 발전시켜 드릴게요.")

    refiner = IdeaRefiner()
    expander = IdeaExpander()
    conv_manager = ConversationManager()

    # 1) 아이디어 입력
    user_idea = input("You: ")

    # 2) 아이디어 정제 (Gemini)
    refined_data = refiner.refine(user_idea)

    print("\nAI: 아이디어를 이렇게 정리해봤어요.")
    print(f"🎯 목표: {refined_data['goal']}")
    print(f"👥 타겟 사용자: {refined_data['target_user']}")
    print("✨ 주요 기능:")
    for feature in refined_data['core_features']:
        print(f"  - {feature}")

    # 3) 확장 아이디어 생성 (Gemini)
    suggestions = expander.expand(refined_data)
    print("\nAI: 이 아이디어를 더 발전시킬 수 있는 몇 가지 제안을 드릴게요.")
    for suggestion in suggestions:
        print(f"- {suggestion}")

    # 4) 유사 기능 검색 (Gemini + GitHub) → map 구성
    print("\nAI: 유사 기능을 검색 중입니다...")
    similar_map = build_similar_map(refined_data['core_features'])
    print("🔍 유사 기능 맵 완성")

    # 5) 대화형 기획 수정
    print("\nAI: 의견을 주시면 반영합니다. (완료/그만/종료 입력 시 종료)")
    final_idea = refined_data
    while True:
        user_feedback = input("You: ")
        if user_feedback in ["완료", "그만", "종료"]:
            break
        final_idea = conv_manager.manage(user_feedback, final_idea)
        print("\nAI: 현재 기획 내용:")
        print(f"🎯 목표: {final_idea['goal']}")
        print(f"👥 타겟 사용자: {final_idea['target_user']}")
        print("✨ 주요 기능:")
        for feature in final_idea['core_features']:
            print(f"  - {feature}")

    # 6) 기획서 생성
    print("\nAI: 기획서 파일을 생성할까요? (네/아니오)")
    if input("You: ").lower() == '네':
        existing_plans = glob.glob("idea_plan_*")
        existing_nums = []
        for f in existing_plans:
            try:
                num_part = f.split('_')[-1].split('.')[0]
                if num_part.isdigit():
                    existing_nums.append(int(num_part))
            except (ValueError, IndexError):
                continue
        next_index = max(existing_nums) + 1 if existing_nums else 1

        md_file = f"idea_plan_{next_index}.md"
        docx_file = f"idea_plan_{next_index}.docx"

        md_content = generate_markdown(final_idea, suggestions, similar_map)
        with open(md_file, "w", encoding="utf-8") as f:
            f.write(md_content)
        generate_word(final_idea, suggestions, similar_map, docx_file)

        print(f"\nAI: 기획서를 `{md_file}`(Markdown)와 `{docx_file}`(Word)로 저장했습니다.")
    else:
        print("AI: 알겠습니다. 필요할 때 다시 호출해주세요.")


if __name__ == "__main__":
    main()
