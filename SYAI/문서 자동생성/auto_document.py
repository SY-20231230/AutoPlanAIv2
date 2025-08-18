#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Auto Document Writer — Interactive + Batch (Gemini 2.5-Flash)
- DOCX/HWPX: {{키}}가 있으면 치환 모드, 없으면 구조-모방 모드(Markdown)
- PDF/HWP : 구조-모방 모드(Markdown)
개선 포인트:
  * [완료] 휴리스틱 실패 시, LLM을 이용한 동적 구조 분석 기능 추가
  * 안전한 LLM 호출(세이프티 완화, 무응답 폴백, 재시도)
  * 헤딩을 배치로 분할하여 섹션 생성(대용량 템플릿 안정화)
  * PDF 긴 텍스트 트림(최대 50KB)
  * 배치별 중간 결과 저장 + 최종 합본 저장
설치:
  pip install google-generativeai python-docx python-dotenv pymupdf
.env:
  GEMINI_API_KEY=YOUR_KEY
"""

import os
import re
import sys
import json
import zipfile
from pathlib import Path
from typing import Dict, List, Tuple, Optional

from dotenv import load_dotenv
import google.generativeai as genai

# ---- Optional imports ----
try:
    import fitz  # PyMuPDF for PDF
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument
    from docx.shared import RGBColor
except Exception:
    DocxDocument = None

# --------------------------
# Gemini 초기화
# --------------------------
def init_gemini(model_name: str):
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY 가 .env 에 없습니다.")
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(model_name)

# --------------------------
# 유틸: 파일/문서 도우미
# --------------------------
PLACEHOLDER_PATTERN = re.compile(r"\{\{([^{}]+)\}\}")

def read_json_or_text(p: Path) -> str:
    raw = p.read_text(encoding="utf-8", errors="ignore")
    try:
        obj = json.loads(raw)
        return json.dumps(obj, ensure_ascii=False, indent=2)
    except Exception:
        return raw

def detect_placeholders_in_text(text: str) -> List[str]:
    return sorted(set(m.group(1).strip() for m in PLACEHOLDER_PATTERN.finditer(text)))

def extract_text_from_pdf(p: Path, max_bytes: int = 50000) -> str:
    if fitz is None:
        return "(PyMuPDF 미설치) PDF 텍스트 추출 불가"
    doc = fitz.open(p.as_posix())
    pages = []
    total = 0
    for i, page in enumerate(doc, start=1):
        if total >= max_bytes:
            break
        try:
            t = page.get_text("text")
        except Exception:
            t = "(텍스트 추출 실패)"
        chunk = f"# Page {i}\n{t}"
        pages.append(chunk)
        total += len(chunk.encode("utf-8", errors="ignore"))
    doc.close()
    return "\n\n".join(pages)

def load_docx_and_plaintext(p: Path) -> Tuple[Optional[object], str]:
    if DocxDocument is None:
        return None, ""
    doc = DocxDocument(p.as_posix())
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            texts.append("\t".join(cell.text for cell in row.cells))
    return doc, "\n".join(texts)

def hwpx_read_xml(p: Path) -> Tuple[bytes, List[str]]:
    with zipfile.ZipFile(p, 'r') as z:
        texts = []
        for name in z.namelist():
            if name.lower().endswith(".xml"):
                with z.open(name) as f:
                    try:
                        data = f.read()
                        texts.append(data.decode("utf-8", errors="ignore"))
                    except Exception:
                        pass
        full = "\n".join(texts)
        placeholders = detect_placeholders_in_text(full)
        return full.encode("utf-8"), placeholders

def hwpx_replace_and_write(src: Path, dst: Path, mapping: Dict[str, str], unsure_to_red=True):
    # 간단 전역치환 방식(실전은 XML 스타일 편집 권장)
    with zipfile.ZipFile(src, 'r') as zin:
        with zipfile.ZipFile(dst, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.lower().endswith(".xml"):
                    text = data.decode("utf-8", errors="ignore")
                    if unsure_to_red:
                        text = re.sub(r"<<UNSURE:(.+?))>>?", r"(불확실: \1)", text)
                    for k, v in mapping.items():
                        text = text.replace(f"{{{{{k}}}}}", v)
                    data = text.encode("utf-8")
                zout.writestr(item, data)

def docx_replace_placeholders(doc, mapping: Dict[str, str], unsure_to_red=True):
    def replace_run_text(run, key, value):
        if f"{{{{{key}}}}}" in run.text:
            run.text = run.text.replace(f"{{{{{key}}}}}", value)

    for para in doc.paragraphs:
        for key, value in mapping.items():
            for run in para.runs:
                replace_run_text(run, key, value)
        if unsure_to_red and "<<UNSURE:" in para.text:
            for run in para.runs:
                m = re.search(r"<<UNSURE:(.+?))>>?", run.text)
                if m:
                    run.text = run.text.replace(m.group(0), m.group(1))
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in mapping.items():
                    cell.text = cell.text.replace(f"{{{{{key}}}}}", value)
                if unsure_to_red and "<<UNSURE:" in cell.text:
                    cell.text = cell.text.replace("<<UNSURE:", "(불확실: ").replace(">>", ")")

# --- MODIFIED FUNCTION ---
# 1. 휴리스틱(규칙 기반) 제목 추출 함수 개선
#    - 더 다양한 패턴을 감지하도록 정규식 확장
#    - 실패 시 빈 리스트를 반환하도록 수정 (하드코딩된 목록 제거)
def extract_headings_heuristic(raw_text: str) -> List[str]:
    """규칙 기반으로 문서에서 제목/항목을 추출합니다. 실패 시 빈 리스트를 반환합니다."""
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]
    headings = []
    # 더 많은 제목 패턴 감지 (예: "1.", "가.", "■", "[1]", "제목:")
    pat = re.compile(
        r"^(\d+(\.\d+)*|[IVXLCM]+\.|[가-힣A-Z]\.|제?\d+장|■|●|□|○|[\(]?\d+[\)]?)\s+|.*:$"
    )
    for ln in lines:
        # 80자 이하의 짧은 줄이면서, 패턴에 맞거나 콜론(:)으로 끝나는 경우
        if len(ln) <= 80 and pat.search(ln):
            # "Page X" 같은 PDF 추출 아티팩트 제거
            if not ln.startswith("# Page"):
                headings.append(ln)

    seen, ordered = set(), []
    for h in headings:
        if h not in seen:
            seen.add(h); ordered.append(h)
    
    # 상위 60개까지만 사용하고, 찾지 못하면 빈 리스트 반환
    return ordered[:60]

# --- NEW FUNCTION ---
# 2. LLM을 이용한 문서 구조(제목) 추출 함수
def extract_structure_with_llm(model, raw_text: str) -> List[str]:
    """휴리스틱으로 제목 찾기 실패 시, LLM을 사용해 문서 구조를 분석하고 제목 목록을 추출합니다."""
    prompt = f"""
다음은 특정 양식에서 추출한 텍스트입니다. 이 문서의 핵심적인 목차, 헤딩, 또는 작성해야 할 항목들을 순서대로 나열해주세요.
- 각 항목은 한 줄에 하나씩만 출력하세요.
- 불필요한 설명이나 번호 매기기는 생략하고, 텍스트에 나타난 제목이나 항목명 자체만 추출하세요.
- 존재하지 않는 항목을 만들지 마세요.
- 만약 분석이 불가능하면, "구조분석불가" 라고만 출력하세요.

[문서 텍스트]
{raw_text[:10000]}
"""
    try:
        resp = model.generate_content(prompt)
        text = resp.text.strip()
        if "구조분석불가" in text or not text:
            return []
        # LLM 응답에서 불필요한 부분을 제거하고 깔끔한 리스트로 변환
        headings = [line.strip() for line in text.splitlines() if line.strip()]
        # 가끔 LLM이 붙이는 '-', '*' 같은 마커 제거
        return [re.sub(r"^[-\*]\s+", "", h) for h in headings if h]
    except Exception as e:
        print(f"[경고] LLM 구조 분석 중 예외 발생: {e}", file=sys.stderr)
        return []

# --------------------------
# 프롬프트 빌더
# --------------------------
def build_prompt(plan_text: str,
                 feature_json_text: Optional[str],
                 task_desc: str,
                 tone: str,
                 style: str,
                 leave_blanks: bool,
                 max_bullets: int = 6,
                 max_chars_per_field: int = 800) -> str:
    rules = f"""
당신은 정부/공모/연구개발 양식에 맞춰 문서를 작성하는 전문 보조자입니다.

[문서 규칙]
- 어투/톤: {tone}
- 문장 스타일: {style} (불릿이면 최대 {max_bullets}개, 항목당 1~2문장)
- 필드/섹션당 최대 {max_chars_per_field}자 이내로 요약
- 과장 금지, 사실 기반 표현. 모호하면 추정하지 마세요.
- 불확실/정보부족: {"빈칸으로 남겨두세요" if leave_blanks else "다음 토큰으로 표시: <<UNSURE: 간단 사유/질문 >>"}
- 개인정보/허위 기재 금지
- 숫자/지표/근거는 가능하면 구체화(없으면 생략)

[참조 기획서/기능명세]
{plan_text}
"""
    if feature_json_text:
        rules += f"\n[추가 기능 명세(JSON)]\n{feature_json_text}\n"
    if task_desc:
        rules += f"\n[작업 설명]\n{task_desc}\n"
    return rules.strip()
# --------------------------
# 프롬프트 & LLM 호출(안전)
# --------------------------
def llm_generate(model, prompt: str, instruction: str) -> str:
    full = prompt + "\n\n[작성 요청]\n" + instruction
    gen_cfg = genai.types.GenerationConfig(
        temperature=0.3, top_p=0.9, max_output_tokens=1536
    )

    # ✅ 버전 호환: dict 형태의 safety_settings 사용
    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUAL", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
    ]

    try:
        resp = model.generate_content(
            full,
            generation_config=gen_cfg,
            safety_settings=safety_settings,
        )
        if getattr(resp, "text", None):
            return resp.text.strip()

        # 폴백: candidates → parts 텍스트 합치기
        for cand in getattr(resp, "candidates", []) or []:
            parts = getattr(getattr(cand, "content", None), "parts", None)
            if parts:
                joined = "".join(getattr(p, "text", "") for p in parts if getattr(p, "text", ""))
                if joined.strip():
                    return joined.strip()

        # 짧게 재시도
        resp2 = model.generate_content(
            full[:6000] + "\n\n(간결히 6줄 이내로.)",
            generation_config=gen_cfg,
            safety_settings=safety_settings,
        )
        if getattr(resp2, "text", None):
            return resp2.text.strip()

    except Exception as e:
        print(f"[경고] LLM 호출 예외: {e}", file=sys.stderr)

    return ""


# --------------------------
# 치환 모드
# --------------------------
def fill_placeholders_with_llm(model,
                               keys: List[str],
                               plan_text: str,
                               feature_json_text: Optional[str],
                               tone: str,
                               style: str,
                               leave_blanks: bool) -> Dict[str, str]:
    mapping = {}
    base_prompt = build_prompt(
        plan_text, feature_json_text,
        "플레이스홀더 항목별로 간결히 작성하세요.",
        tone, style, leave_blanks
    )
    for k in keys:
        instruction = f"""다음 필드를 작성하세요.

[필드명]
{k}

[출력 형식]
- 순수 텍스트만.
- {("빈칸 허용" if leave_blanks else "불확실하면 <<UNSURE: ...>> 로 표시")}
"""
        txt = llm_generate(model, base_prompt, instruction)
        if leave_blanks and (txt.strip().lower() in {"n/a", "없음", "불명", "미상"} or "unsure" in txt.lower()):
            txt = ""
        mapping[k] = txt
    return mapping

# --------------------------
# 구조-모방 모드 (배치)
# --------------------------
def chunk_list(items: List[str], size: int) -> List[List[str]]:
    return [items[i:i+size] for i in range(0, len(items), size)]

def mimic_structure_with_llm_batch(model,
                                   headings: List[str],
                                   plan_text: str,
                                   feature_json_text: Optional[str],
                                   tone: str,
                                   style: str,
                                   leave_blanks: bool,
                                   batch_size: int,
                                   stem_for_save: Path) -> str:
    """헤딩을 batch_size로 나눠 섹션 생성, 배치별 중간파일 저장, 전체 합본 반환."""
    base_prompt = build_prompt(
        plan_text, feature_json_text,
        "문서의 각 섹션을 해당 헤딩 아래에 작성하세요.",
        tone, style, leave_blanks
    )
    all_sections: List[str] = []
    batches = chunk_list(headings, max(1, batch_size))
    for bi, batch in enumerate(batches, start=1):
        print(f"  - 배치 {bi}/{len(batches)} (헤딩 {len(batch)}개) 생성 중...")
        batch_blocks = []
        for h in batch:
            instruction = f"""아래 헤딩 섹션을 작성하세요.

[헤딩]
{h}

[출력 형식]
- Markdown 섹션으로 출력 (예: '## {h}')
- 본문은 {style} 기준
- {("빈칸 허용" if leave_blanks else "불확실시 <<UNSURE: ...>>")}
"""
            out = llm_generate(model, base_prompt, instruction)
            if not out.strip():
                continue  # 빈 응답은 스킵
            if not out.strip().startswith("#"):
                out = f"## {h}\n\n{out}"
            batch_blocks.append(out)
            all_sections.append(out)

        # 배치별 중간 저장
        if batch_blocks:
            md_part = "\n\n".join(batch_blocks)
            part_path = stem_for_save.with_name(f"{stem_for_save.stem}_draft_part{bi}.md")
            part_path.write_text(md_part, encoding="utf-8")
            print(f"    · 저장: {part_path}")

    # 합본 저장은 호출부에서 처리
    return "\n\n".join(all_sections)

# --------------------------
# 보조: 출력 경로 보정
# --------------------------
def resolve_output_path(user_in: str, default_path: Path, desired_suffix: str) -> Path:
    """사용자가 폴더만 줬거나 확장자가 없으면 자동 보정."""
    if not user_in:
        return default_path
    p = Path(user_in)
    if p.is_dir():
        return p / (default_path.name)
    if p.suffix == "":
        return p.with_suffix(desired_suffix)
    # 폴더가 아직 없으면 생성
    if not p.parent.exists():
        try:
            p.parent.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass
    return p

# --- MODIFIED HELPER FUNCTION ---
# 3. 문서 처리 로직을 별도 함수로 분리 (main 단순화)
def process_structure_mimic_mode(model, raw_text: str, tpl_path: Path, plan_text: str, feature_json_text: Optional[str], tone: str, style: str, leave_blanks: bool, batch_size: int, out_override: str):
    """구조-모방 모드를 실행하고, 필요시 LLM으로 구조를 분석합니다."""
    # 1. 휴리스틱으로 제목 추출 시도
    headings = extract_headings_heuristic(raw_text)

    # 2. 실패 시, LLM으로 구조 분석
    if not headings:
        print("  - 휴리스틱 제목 탐지 실패. LLM으로 문서 구조를 분석합니다...")
        headings = extract_structure_with_llm(model, raw_text)
    
    if not headings:
        print("[오류] 문서에서 작성할 항목(헤딩)을 추출할 수 없습니다. 문서 내용을 확인해주세요.", file=sys.stderr)
        # LLM 분석도 실패하면, 최소한의 기본 목차로 동작하도록 폴백
        headings = [
            "1. 개요 (Introduction)",
            "2. 주요 내용 (Main Content)",
            "3. 기대 효과 (Expected Outcomes)",
            "4. 결론 (Conclusion)"
        ]
        print("  - 기본 목차로 초안 생성을 시도합니다:", headings)

    print(f"  - 총 {len(headings)}개의 헤딩을 기반으로 문서 생성을 시작합니다.")

    stem_for_save = tpl_path.with_name(tpl_path.stem)
    md_all = mimic_structure_with_llm_batch(model, headings, plan_text, feature_json_text, tone, style, leave_blanks, batch_size, stem_for_save)
    default_out = tpl_path.with_name(tpl_path.stem + "_draft.md")
    outp = resolve_output_path(out_override, default_out, ".md")
    outp.write_text(md_all, encoding="utf-8")
    print(f"✅ 합본 저장: {outp}")

    if DocxDocument:
        mk_docx = input("➡️ 초안을 DOCX로도 저장할까요? (y/N): ").strip().lower() == "y"
        if mk_docx:
            d2 = DocxDocument()
            # 마크다운 헤딩을 좀 더 보기 좋게 변환
            for block in md_all.split('\n\n'):
                lines = block.split('\n')
                if lines[0].startswith('## '):
                    d2.add_heading(lines[0][3:].strip(), level=2)
                    d2.add_paragraph('\n'.join(lines[1:]))
                else:
                    d2.add_paragraph(block)
                d2.add_paragraph() # 문단 간격

            d2_out = outp.with_suffix(".docx")
            d2.save(d2_out.as_posix())
            print(f"(변환) DOCX: {d2_out}")

# --------------------------
# 인터랙티브 메인
# --------------------------
def main():
    print("📄 Auto Document Writer (Gemini 2.5-Flash) — Batch\n")

    # 1) 입력 받기
    plan_path_str = input("📝 기획서/기능명세 파일 경로 (.json 또는 .txt): ").strip()
    template_path_str = input("📂 작성할 양식 파일 경로 (.hwpx/.docx/.pdf/.hwp): ").strip()

    tone = input("💬 어투/톤 (기본: 격식 있는 비즈니스 한국어): ").strip() or "기본"
    style = input("✏️ 문장 스타일 (기본: 개조식 불릿): ").strip() or "개조식"
    leave_blanks_input = input("❓ 불확실시 빈칸 처리? (y/N): ").strip().lower()
    leave_blanks = (leave_blanks_input == "y")
    model_name = input("🤖 Gemini 모델명 (기본: gemini-2.5-flash): ").strip() or "gemini-2.5-flash"
    if model_name in {"기본", "default", "Default"}:
        model_name = "gemini-2.5-flash"

    try:
        batch_size = int(input("📦 배치 크기(한 번에 생성할 헤딩 수, 기본 10): ").strip() or "10")
    except ValueError:
        batch_size = 10

    out_override = input("💾 출력 파일/폴더 경로(미입력 시 자동 결정): ").strip()

    plan_path = Path(plan_path_str)
    tpl_path = Path(template_path_str)

    if not plan_path.exists():
        print("기획서 파일을 찾을 수 없습니다.", file=sys.stderr); sys.exit(1)
    if not tpl_path.exists():
        print("양식 파일을 찾을 수 없습니다.", file=sys.stderr); sys.exit(1)

    # 2) 모델 준비
    try:
        model = init_gemini(model_name)
    except Exception as e:
        print(f"Gemini 초기화 실패: {e}", file=sys.stderr)
        sys.exit(1)

    # 3) 기획서 읽기
    plan_text = read_json_or_text(plan_path)
    feature_json_text = None
    try:
        obj = json.loads(plan_text)
        feature_json_text = json.dumps(obj, ensure_ascii=False, indent=2)
    except Exception:
        pass

    # 4) 분기 처리
    ext = tpl_path.suffix.lower()

    if ext == ".docx":
        if DocxDocument is None:
            print("python-docx 미설치로 .docx 처리가 불가합니다.", file=sys.stderr); sys.exit(1)
        doc, plain = load_docx_and_plaintext(tpl_path)
        keys = detect_placeholders_in_text(plain)
        if keys:
            print(f"[치환 모드] 감지된 키: {keys}")
            mapping = fill_placeholders_with_llm(model, keys, plan_text, feature_json_text, tone, style, leave_blanks)
            docx_replace_placeholders(doc, mapping, unsure_to_red=(not leave_blanks))
            default_out = tpl_path.with_name(tpl_path.stem + "_filled.docx")
            outp = resolve_output_path(out_override, default_out, ".docx")
            doc.save(outp.as_posix())
            print(f"✅ 완료: {outp}")
        else:
            print("[구조-모방 모드] DOCX 텍스트 기반으로 섹션 생성(배치)")
            process_structure_mimic_mode(model, plain, tpl_path, plan_text, feature_json_text, tone, style, leave_blanks, batch_size, out_override)

    elif ext == ".hwpx":
        xml_blob, keys = hwpx_read_xml(tpl_path)
        if keys:
            print(f"[치환 모드] HWPX 키 감지: {keys}")
            mapping = fill_placeholders_with_llm(model, keys, plan_text, feature_json_text, tone, style, leave_blanks)
            default_out = tpl_path.with_name(tpl_path.stem + "_filled.hwpx")
            outp = resolve_output_path(out_override, default_out, ".hwpx")
            hwpx_replace_and_write(tpl_path, outp, mapping, unsure_to_red=(not leave_blanks))
            print(f"✅ 완료: {outp}")
        else:
            print("[구조-모방 모드] HWPX → Markdown 초안(배치)")
            raw_text = xml_blob.decode("utf-8", errors="ignore")
            process_structure_mimic_mode(model, raw_text, tpl_path, plan_text, feature_json_text, tone, style, leave_blanks, batch_size, out_override)

    elif ext == ".pdf":
        print("[구조-모방 모드] PDF 텍스트 추출(최대 50KB) → Markdown 초안(배치)")
        raw = extract_text_from_pdf(tpl_path, max_bytes=50000)
        process_structure_mimic_mode(model, raw, tpl_path, plan_text, feature_json_text, tone, style, leave_blanks, batch_size, out_override)

    elif ext == ".hwp":
        print("[구조-모방 모드] HWP는 구조 추출 제한 → Markdown 초안(배치)")
        raw = "(HWP 텍스트 추출 미구현: HWPX 변환 권장)"
        # HWP의 경우 텍스트 추출이 안되므로 LLM 분석도 의미가 없음.
        # 이 경우에는 정말 최소한의 목차만 제공
        headings = [
            "1. 과제 개요",
            "2. 주요 내용 및 추진 전략",
            "3. 기대효과 및 활용방안"
        ]
        print("  - HWP 파일은 내용 분석이 불가하여 기본 목차로 생성합니다.")
        stem_for_save = tpl_path.with_name(tpl_path.stem)
        md_all = mimic_structure_with_llm_batch(model, headings, plan_text, feature_json_text, tone, style, leave_blanks, batch_size, stem_for_save)
        default_out = tpl_path.with_name(tpl_path.stem + "_draft.md")
        outp = resolve_output_path(out_override, default_out, ".md")
        outp.write_text(md_all, encoding="utf-8")
        print(f"✅ 합본 저장: {outp}")

    else:
        print(f"지원하지 않는 확장자: {ext}", file=sys.stderr)
        sys.exit(2)

if __name__ == "__main__":
    main()