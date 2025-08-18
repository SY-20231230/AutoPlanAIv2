from typing import Optional
from docx import Document


def new_doc() -> Document:
	return Document()


def append_markdownish_section(doc: Document, title: str, md_text: str) -> None:
	# 섹션 제목
	doc.add_heading(title, level=2)
	if not md_text:
		return
	for raw in md_text.splitlines():
		line = raw.rstrip()
		if not line:
			# 빈 줄은 단락 구분
			doc.add_paragraph("")
			continue
		if line.startswith("### "):
			doc.add_heading(line[4:].strip(), level=3)
			continue
		if line.startswith("## "):
			doc.add_heading(line[3:].strip(), level=3)
			continue
		if line.startswith("- "):
			p = doc.add_paragraph(style="List Bullet")
			p.add_run(line[2:].strip())
			continue
		# 기본 단락
		doc.add_paragraph(line)


def save_doc(doc: Document, path: str) -> None:
	doc.save(path)
